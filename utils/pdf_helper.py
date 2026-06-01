"""
utils/pdf_helper.py
Converts PDF / DOCX / plain text into a list of clause dicts.
Returns: [{"id": "clause_1", "text": "..."}, ...]
"""
import re
import io

def extract_clauses_from_pdf(file_bytes: bytes) -> list:
    import PyPDF2
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    raw = ""
    for page in reader.pages:
        raw += (page.extract_text() or "") + "\n"
    return _split_into_clauses(raw)

def extract_clauses_from_docx(file_bytes: bytes) -> list:
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))
    raw = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return _split_into_clauses(raw)

def extract_clauses_from_text(text: str) -> list:
    return _split_into_clauses(text)

def extract_text_for_summary(file_bytes: bytes, filetype: str) -> str:
    """For document summarizer — returns raw text (first 2500 chars)."""
    if filetype == "pdf":
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text = " ".join(p.extract_text() or "" for p in reader.pages[:3])
    elif filetype == "docx":
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        text = file_bytes.decode("utf-8", errors="ignore")
    return text[:2500]

def _split_into_clauses(text: str) -> list:
    """
    Split contract text into individual clauses.
    Handles: numbered (1. 2. 3.), lettered (a. b.), and paragraph splits.
    Each clause becomes one item for AI analysis.
    """
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Split on numbered clause patterns: "1." "2." "(1)" "(a)" at start of line or after period+space
    # This regex splits BEFORE each numbered item
    split_pattern = re.compile(
        r'(?<!\w)'           # not preceded by word char
        r'(?='               # lookahead
        r'(?:\d{1,2}\.\ )'  # "1. " "2. " etc
        r'|(?:\([a-zA-Z0-9]\)\ )'  # "(a) " "(1) " etc
        r')',
        re.MULTILINE
    )

    parts = split_pattern.split(text)

    clauses = []
    clause_id = 1

    for part in parts:
        part = part.strip()
        if not part or len(part) < 20:
            continue

        # If the part is very long (>300 chars), split at sentence boundaries
        if len(part) > 300:
            sentences = re.split(r'(?<=[.!?])\s+', part)
            buffer = ""
            for sentence in sentences:
                buffer += (" " if buffer else "") + sentence
                if len(buffer) > 200:
                    clauses.append({
                        "id": f"clause_{clause_id}",
                        "text": buffer.strip()
                    })
                    clause_id += 1
                    buffer = ""
            if buffer.strip() and len(buffer.strip()) > 20:
                clauses.append({
                    "id": f"clause_{clause_id}",
                    "text": buffer.strip()
                })
                clause_id += 1
        else:
            clauses.append({
                "id": f"clause_{clause_id}",
                "text": part
            })
            clause_id += 1

    return clauses