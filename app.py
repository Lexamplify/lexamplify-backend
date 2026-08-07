from flask import Flask, render_template, jsonify, request, redirect
from flask_cors import CORS
from flask_jwt_extended import JWTManager
from dotenv import load_dotenv
from database import db as sqlalchemy_db
from datetime import timedelta
import os
import re
import secrets
import sqlite3
import json
import io
import time
import uuid
import requests
import urllib.parse
from bs4 import BeautifulSoup
from groq import Groq
from tavily import TavilyClient
from pinecone import Pinecone
from utils.ai_helper import extract_json_from_llm_response

# Render's filesystem is ephemeral — anything written to the working
# directory (including a bare-relative-path SQLite file like
# lex_assistant.db) is wiped on every deploy/restart unless it lives on an
# attached Persistent Disk. Every sqlite3.connect('lex_assistant.db') call
# across this codebase (10+ files: app.py, document_routes.py,
# auth_routes.py, ai_routes.py, conflict_routes.py, rag_pipeline.py, ...)
# uses that same bare relative path, resolved against the process's CWD —
# rather than threading an env-var-driven absolute path through every one
# of those call sites, this symlinks the CWD-relative name to the
# Persistent Disk's mount path ONCE, at import time, before anything
# connects. Every existing sqlite3.connect('lex_assistant.db') call then
# transparently resolves through the symlink with zero further changes.
# PERSISTENT_DATA_DIR is only set on Render once a Disk is attached and
# this env var points at its mount path — locally it's unset, so this
# block is a no-op and lex_assistant.db resolves exactly as it always has.
load_dotenv()
_persistent_data_dir = os.getenv('PERSISTENT_DATA_DIR')
if _persistent_data_dir and os.path.isdir(_persistent_data_dir):
    _persistent_db_path = os.path.join(_persistent_data_dir, 'lex_assistant.db')
    if not os.path.exists(_persistent_db_path):
        # First boot after attaching a fresh disk — nothing to link to yet;
        # sqlite3 will populate real tables into this empty file once
        # init_db()/init_sqlite_db() run against it through the symlink.
        open(_persistent_db_path, 'a').close()
    if not os.path.islink('lex_assistant.db'):
        if os.path.exists('lex_assistant.db'):
            # A non-symlink file already exists here (e.g. this is the very
            # first deploy after attaching the disk) — never silently
            # discard it; back it up once instead of overwriting.
            os.rename('lex_assistant.db', 'lex_assistant.db.pre-disk-backup')
        os.symlink(_persistent_db_path, 'lex_assistant.db')

db = sqlite3.connect('lex_assistant.db', check_same_thread=False)

def init_db():
    conn = db
    c = conn.cursor()
    c.execute('DROP TABLE IF EXISTS calendar_events')
    c.execute('''
        CREATE TABLE IF NOT EXISTS calendar_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            event_date TEXT,
            event_type TEXT,
            title TEXT,
            related_case_id TEXT,
            location TEXT,
            opposing_counsel TEXT
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS case_vault (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            case_id TEXT,
            title TEXT,
            doc_type TEXT,
            content TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    # Adjacency-list folder hierarchy for Case Vault
    c.execute('''
        CREATE TABLE IF NOT EXISTS vault_folders (
            id   INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT    NOT NULL,
            parent_id INTEGER REFERENCES vault_folders(id) ON DELETE CASCADE,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    # Migration: add folder_id + smart_title + tags + file_blob + file_format to case_vault if not present
    for _col, _type in (('folder_id', 'INTEGER'), ('smart_title', 'TEXT'), ('tags', 'TEXT'), ('file_blob', 'BLOB'), ('file_format', 'TEXT')):
        try:
            c.execute(f'ALTER TABLE case_vault ADD COLUMN {_col} {_type}')
        except Exception:
            pass
    # AI Provenance / Audit Trail — links saved docs to the agent conversation that produced them
    c.execute('''
        CREATE TABLE IF NOT EXISTS vault_audit (
            id             INTEGER PRIMARY KEY AUTOINCREMENT,
            folder_id      INTEGER REFERENCES vault_folders(id) ON DELETE SET NULL,
            vault_doc_id   INTEGER,
            session_title  TEXT,
            messages_json  TEXT,
            created_at     DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS document_chunks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            case_id INTEGER,
            document_id INTEGER NOT NULL,
            chunk_index INTEGER NOT NULL,
            chunk_text TEXT NOT NULL,
            embedding TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    c.execute('''
        CREATE TABLE IF NOT EXISTS chat_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            user_message TEXT NOT NULL,
            bot_response TEXT NOT NULL,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.commit()

load_dotenv()

jwt = JWTManager()

DB_PATH = os.path.join('instance', 'client_data.db')

def init_sqlite_db():
    os.makedirs('instance', exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS tracked_cases (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        case_name TEXT NOT NULL,
        case_number TEXT,
        cnr_number TEXT,
        court TEXT,
        client_name TEXT,
        case_type TEXT,
        next_hearing DATE,
        last_hearing DATE,
        status TEXT DEFAULT 'Active',
        notes TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        cnr TEXT,
        title TEXT,
        next_hearing_date DATE
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS clients (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        phone TEXT,
        email TEXT,
        address TEXT,
        client_type TEXT DEFAULT 'Individual',
        notes TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        contact TEXT,
        company TEXT
    )''')
    
    # Ensure missing schema columns are dynamically appended
    for col, col_type in [("cnr", "TEXT"), ("title", "TEXT"), ("next_hearing_date", "DATE")]:
        try:
            c.execute(f"ALTER TABLE tracked_cases ADD COLUMN {col} {col_type}")
        except sqlite3.OperationalError:
            pass
            
    for col, col_type in [("contact", "TEXT"), ("company", "TEXT")]:
        try:
            c.execute(f"ALTER TABLE clients ADD COLUMN {col} {col_type}")
        except sqlite3.OperationalError:
            pass
    c.execute('''CREATE TABLE IF NOT EXISTS time_entries (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        client_id INTEGER,
        case_id INTEGER,
        date DATE,
        hours REAL,
        rate REAL,
        description TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS invoices (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        client_id INTEGER,
        invoice_number TEXT,
        amount REAL,
        gst REAL,
        total REAL,
        status TEXT DEFAULT 'Draft',
        due_date DATE,
        paid_date DATE,
        notes TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS ip_assets (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ip_type TEXT,
        title TEXT,
        registration_number TEXT,
        filing_date DATE,
        renewal_due DATE,
        status TEXT DEFAULT 'Active',
        notes TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS team_members (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        email TEXT,
        role TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS tasks (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT,
        assigned_to INTEGER,
        case_id INTEGER,
        due_date DATE,
        priority TEXT DEFAULT 'Normal',
        status TEXT DEFAULT 'todo',
        notes TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    conn.commit()
    conn.close()


def generate_pdf_blob(title, content):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm,
                            leftMargin=2.5*cm, rightMargin=2.5*cm)
    styles = getSampleStyleSheet()
    story = [Paragraph(str(title), styles['Title']), Spacer(1, 0.4*cm)]
    for line in str(content).split('\n'):
        stripped = line.strip()
        if stripped:
            story.append(Paragraph(stripped, styles['Normal']))
            story.append(Spacer(1, 0.15*cm))
        else:
            story.append(Spacer(1, 0.25*cm))
    doc.build(story)
    return buf.getvalue()


def generate_docx_blob(title, content):
    from docx import Document
    buf = io.BytesIO()
    doc = Document()
    doc.add_heading(str(title), level=1)
    for line in str(content).split('\n'):
        doc.add_paragraph(line)
    doc.save(buf)
    return buf.getvalue()


# Matches any http(s) localhost/127.0.0.1 origin regardless of port, so a
# Vite dev-server port bump (5173 -> 5174 -> ...) never silently breaks CORS.
# Anchored start-to-end so it can't be satisfied by a lookalike host like
# "http://localhost.evil.com:5173".
LOCAL_DEV_ORIGIN_RE = re.compile(r"^https?://(localhost|127\.0\.0\.1):\d+$")
PROD_ORIGINS = {
    'https://lexamplify-4.web.app',
    'https://lexamplify-4.firebaseapp.com',
    'https://test.lexamplify.com',
    'https://lexamplify.com',
}

# Strict allow-list for the document-retrieval route's case_id path segment —
# alphanumeric plus underscore/hyphen only, anchored start-to-end, so a
# value like "../../etc/passwd" can never reach the SQL query.
CASE_ID_RE = re.compile(r'^[A-Za-z0-9_-]+$')

# worker.py's Pinecone vector IDs are hashlib.md5(...).hexdigest() of
# "{case_id}_chunk_{n}" — opaque 32-char hashes that don't decode back to a
# case_id. A URL path segment matching this shape means the caller only had
# the raw vector id (e.g. a metadata-less search hit), not the real case_id.
MD5_HASH_RE = re.compile(r'^[A-Za-z0-9]{32}$')

# Same strict allow-list intent as CASE_ID_RE, but worker.py's real case_ids
# are raw S3 filenames like "2022_13_342_356_EN.pdf" — CASE_ID_RE (built for
# the case_vault route's slug-style ids) has no '.', which would reject
# every single legitimately-ingested judgment. Still anchored start-to-end
# and still excludes '/', quotes, and shell/SQL metacharacters.
PINECONE_CASE_ID_RE = re.compile(r'^[A-Za-z0-9_.-]+$')

# Must exactly match the ingestion pipeline's embedding model/index
# dimension (worker.py EMBED_MODEL = "llama-text-embed-v2") — verified
# empirically against the live index via describe_index_stats(): dimension
# 1024, metric cosine. A true zero vector is undefined under cosine
# similarity (division by zero); this dummy is only ever used for
# filter-only "give me everything matching this case_id" queries where the
# similarity score itself is discarded, so its direction is irrelevant —
# it only needs to be non-zero.
EMBED_DIMENSION = 1024
EMBED_MODEL = "llama-text-embed-v2"
DUMMY_QUERY_VECTOR = [1e-5] * EMBED_DIMENSION


def format_case_title(filename):
    """Builds a display title from a worker.py-style case_id, e.g.
    "2022_13_342_356_EN.pdf" -> "Supreme Court Judgment (2022) - 2022_13_342_356".
    Null-safe: the caller may pass metadata straight from Pinecone, which is
    pipeline-controlled and not guaranteed to have every field.
    """
    if not filename:
        return "Supreme Court Judgment (Unknown) - Untitled"

    clean_name = filename
    if clean_name.endswith('_EN.pdf'):
        clean_name = clean_name[:-len('_EN.pdf')]
    elif clean_name.endswith('.pdf'):
        clean_name = clean_name[:-len('.pdf')]

    parts = clean_name.split('_')
    year = parts[0] if parts and parts[0] else "Unknown"
    return f"Supreme Court Judgment ({year}) - {clean_name}"


def _get_pinecone_index(pc):
    """Resolves the Pinecone index by host if PINECONE_HOST is configured
    (adding the https:// scheme if the env var was saved without one),
    else by name — shared by every route that queries the case-law index
    so a malformed/missing env var can't crash it differently in two
    places. Callers are expected to wrap this in their own try/except;
    it deliberately does not swallow errors itself."""
    raw_host = (os.getenv("PINECONE_HOST") or "").strip()
    index_name = (os.getenv("PINECONE_INDEX_NAME") or "legal-cases").strip()
    if raw_host:
        if not raw_host.startswith("https://"):
            raw_host = f"https://{raw_host}"
        return pc.Index(host=raw_host)
    return pc.Index(name=index_name)


# Distinct from format_case_title above — that one parses worker.py's S3
# filename-style ids ("2022_13_342_356_EN.pdf"). This is for case_vault
# rows, which use a different id scheme entirely (bulk INSC neutral
# citations like "2025INSC738", or hand-entered ids like "case_101").
# Confirmed empirically against the live table: title is NEVER null/empty,
# but 894 of 911 rows have title == case_id verbatim (a bulk-ingest job set
# title to a bare copy of the docket id instead of a real case name) — so
# "does a title exist" is the wrong check; "is the title actually different
# information from the id" is the one that matches real data.
INSC_CITATION_RE = re.compile(r'^(\d{4})INSC(\d+)$', re.IGNORECASE)


def format_vault_title(raw_title, case_id):
    candidate = (raw_title or '').strip()
    fallback_id = (case_id or '').strip()
    if candidate and candidate != fallback_id:
        return candidate

    source = candidate or fallback_id
    if not source:
        return 'Untitled Document'

    m = INSC_CITATION_RE.match(source)
    if m:
        year, num = m.groups()
        return f"{year} INSC {num}"
    return source


# Party-name extraction for case_vault rows whose title is uninformative
# (title == case_id). Scans line-by-line rather than running the regex
# against the whole document: judgments in this table can be hundreds of
# KB, and a pattern like ".*v\..*" applied with DOTALL/re.search over the
# full text risks catastrophic backtracking and would happily match across
# paragraph boundaries, picking up garbage. Anchoring to one line at a time
# keeps each match bounded to at most ~120 chars of work.
PARTY_VS_LINE_RE = re.compile(r"([^\n]{3,60}\s+(?:v\.|vs\.|vs|VERSUS)\s+[^\n]{3,60})", re.IGNORECASE)
STANDALONE_VS_RE = re.compile(r'^(?:v\.|vs\.|vs|versus)$', re.IGNORECASE)

# Case captions live in the first few lines of a judgment; a "Case Law
# Cited" section further down is full of other "Party v. Party" lines
# (precedents, not the case's own parties) and must never win.
CAPTION_SCAN_WINDOW = 60


def extract_case_title_from_content(content):
    if not content:
        return None
    lines = [l.strip() for l in content.split('\n')]
    non_empty = [l for l in lines if l][:CAPTION_SCAN_WINDOW]
    for i, line in enumerate(non_empty):
        m = PARTY_VS_LINE_RE.search(line)
        if m:
            cleaned = re.sub(r'\s+', ' ', m.group(1)).strip(' .,-:;')
            if cleaned:
                return cleaned
        # Official SCR-style captions often put the separator on its own
        # line ("Party A\nv.\nParty B") — join it with its immediate
        # neighbours rather than a document-wide greedy/DOTALL match.
        if STANDALONE_VS_RE.match(line) and 0 < i < len(non_empty) - 1:
            party_a, party_b = non_empty[i - 1], non_empty[i + 1]
            if 3 <= len(party_a) <= 80 and 3 <= len(party_b) <= 80:
                cleaned = re.sub(r'\s+', ' ', f'{party_a} {line} {party_b}').strip(' .,-:;')
                if cleaned:
                    return cleaned
    return None


def resolve_vault_title(raw_title, case_id, content):
    """Best available title for a case_vault row: a genuinely descriptive
    stored title, else a 'Party v. Party' line extracted from the document
    body, else the INSC-citation-formatted case_id (format_vault_title)."""
    candidate = (raw_title or '').strip()
    fallback_id = (case_id or '').strip()
    if candidate and candidate != fallback_id:
        return candidate
    extracted = extract_case_title_from_content(content)
    if extracted:
        return extracted
    return format_vault_title(raw_title, case_id)


def create_app():
    app = Flask(__name__)
    @app.after_request
    def add_cors_headers(response):
        origin = request.headers.get('Origin')
        if origin and (origin in PROD_ORIGINS or LOCAL_DEV_ORIGIN_RE.match(origin)):
            # Echo back the validated origin rather than hardcoding one — never
            # assign '*', since Allow-Credentials requires a specific origin.
            # Use direct assignment so we never produce duplicate CORS headers
            # (SSE responses pre-set these; .add() would append a second value)
            response.headers['Access-Control-Allow-Origin'] = origin
            response.headers['Access-Control-Allow-Headers'] = 'Content-Type,Authorization,Accept,X-Requested-With,X-CSRF-TOKEN'
            response.headers['Access-Control-Allow-Methods'] = 'GET,PUT,POST,DELETE,OPTIONS'
            response.headers['Access-Control-Allow-Credentials'] = 'true'
        return response

    app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///database.db')
    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

    # ── JWT secret: no static fallback, ever ─────────────────────────────
    # A hardcoded 'secret' default meant every unconfigured deployment
    # shared the same signing key. In production, an unset key is a hard
    # failure — fail fast at boot, not silently sign tokens with a guessable
    # key. In dev, generate a real 32-byte key per process so
    # InsecureKeyLengthWarning can't fire even when .env is untouched;
    # restarting invalidates existing sessions, which is acceptable for a
    # local dev loop and not something production should ever do.
    _jwt_secret = os.getenv('JWT_SECRET_KEY')
    if not _jwt_secret:
        if os.getenv('FLASK_ENV') == 'production':
            raise RuntimeError('JWT_SECRET_KEY must be set in production — no fallback is generated.')
        _jwt_secret = secrets.token_hex(32)
        print('[app] JWT_SECRET_KEY not set — generated an ephemeral dev key (sessions will not survive a restart).')
    app.config['JWT_SECRET_KEY'] = _jwt_secret

    app.config['UPLOAD_FOLDER'] = os.getenv('UPLOAD_FOLDER', 'static/uploads')
    app.config['MAX_CONTENT_LENGTH'] = int(os.getenv('MAX_CONTENT_LENGTH', 104857600))  # 100 MB
    app.config['SQLITE_DB_PATH'] = DB_PATH

    # ── JWT: HttpOnly cookie transport, never localStorage ──────────────
    # Short-lived access token + long-lived refresh token is what makes
    # this "persistent" without keeping a broadly-scoped token alive for
    # 30 days — /api/auth/refresh silently rotates a new access cookie
    # from the refresh cookie, so the session survives reloads without the
    # frontend ever holding the JWT itself in JS-reachable storage.
    app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(hours=1)
    app.config['JWT_REFRESH_TOKEN_EXPIRES'] = timedelta(days=30)
    app.config['JWT_TOKEN_LOCATION'] = ['cookies', 'headers']  # 'headers' kept only for any not-yet-migrated caller; cookies are primary
    app.config['JWT_COOKIE_HTTPONLY'] = True
    app.config['JWT_COOKIE_SECURE'] = os.getenv('JWT_COOKIE_SECURE', 'False').lower() == 'true'
    # Strict now works end-to-end in dev because frontend/vite.config.js
    # proxies /api to Flask — the browser only ever talks to the Vite
    # origin, so the cookie is first-party from its point of view. A
    # production deploy needs the equivalent same-origin proxy/rewrite
    # (e.g. a reverse proxy or platform rewrite rule) in front of the API
    # for this to keep working; without one, Strict drops the cookie on
    # every request the way it did before the proxy existed.
    app.config['JWT_COOKIE_SAMESITE'] = 'Strict'
    app.config['JWT_COOKIE_CSRF_PROTECT'] = True
    app.config['JWT_CSRF_IN_COOKIES'] = True  # readable (non-HttpOnly) CSRF cookie for the double-submit header

    # Flask's OWN session cookie (distinct from the JWT cookies above) —
    # used only by Authlib to hold OAuth `state` across the redirect to
    # Microsoft/Google and back. Must be Lax, not Strict: the IdP's
    # callback redirect is itself a cross-site top-level navigation, and a
    # Strict session cookie would not be sent on it, breaking SSO outright.
    app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
    app.config['SESSION_COOKIE_SECURE'] = app.config['JWT_COOKIE_SECURE']
    app.config['SESSION_COOKIE_HTTPONLY'] = True

    # ── Celery (Redis broker/backend) ────────────────────────────────────
    # Docker provisions ONLY Redis (docker-compose.yml) — the Celery worker
    # runs locally via `celery -A celery_worker.celery worker`, sharing this
    # same venv/dependencies. Registering the extension here (not calling
    # celery_init_app itself, to avoid a celery_app<->app.py import cycle at
    # module load time) is what lets routes reach current_app.extensions
    # 'celery' without re-importing app.py from inside a task module.
    app.config['CELERY_BROKER_URL'] = os.getenv('CELERY_BROKER_URL', 'redis://localhost:6379/0')
    app.config['CELERY_RESULT_BACKEND'] = os.getenv('CELERY_RESULT_BACKEND', 'redis://localhost:6379/0')
    from celery_app import celery_init_app
    celery_init_app(app)

    sqlalchemy_db.init_app(app)
    jwt.init_app(app)

    from utils.limiter import limiter
    limiter.init_app(app)

    from routes.sso_routes import init_sso
    init_sso(app)

    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    init_sqlite_db()

    @app.route('/api/ai/extract-file', methods=['POST', 'OPTIONS'])
    def extract_file_text():
        """Extract plain text from an uploaded PDF, DOCX, or TXT for the AI Legal Associate."""
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        f = request.files.get('file')
        if not f or not f.filename:
            return jsonify({'error': 'No file provided'}), 400
        ext = f.filename.rsplit('.', 1)[-1].lower() if '.' in f.filename else 'txt'
        try:
            from utils.pdf_helper import extract_text_for_summary
            raw_bytes = f.read()
            text = extract_text_for_summary(raw_bytes, ext)
            if not text or not text.strip():
                text = raw_bytes.decode('utf-8', errors='ignore')
            return jsonify({'text': text[:12000], 'filename': f.filename})
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    # ── Vault Folders API ─────────────────────────────────────────────────────
    @app.route('/api/vault/folders', methods=['GET', 'OPTIONS'])
    def get_vault_folders():
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            conn = db
            old_rf = conn.row_factory
            conn.row_factory = sqlite3.Row
            try:
                rows = conn.execute(
                    'SELECT id, name, parent_id, created_at FROM vault_folders ORDER BY name ASC'
                ).fetchall()
                folders = [dict(r) for r in rows]

                # Cheap GROUP BY for per-folder document counts — never touches
                # `content`, so this stays fast regardless of table size and
                # lets the frontend show folder counts without loading every
                # document just to count them client-side.
                count_rows = conn.execute(
                    'SELECT folder_id, COUNT(*) AS cnt FROM case_vault GROUP BY folder_id'
                ).fetchall()
                doc_counts = {
                    ('root' if r['folder_id'] is None else str(r['folder_id'])): r['cnt']
                    for r in count_rows
                }
            finally:
                conn.row_factory = old_rf

            # Build adjacency list → nested tree
            by_id = {f['id']: {**f, 'children': []} for f in folders}
            roots = []
            for f in by_id.values():
                pid = f.get('parent_id')
                if pid and pid in by_id:
                    by_id[pid]['children'].append(f)
                else:
                    roots.append(f)

            return jsonify({'folders': roots, 'flat': folders, 'doc_counts': doc_counts}), 200
        except Exception as e:
            return jsonify({'error': True, 'message': str(e)}), 500

    @app.route('/api/vault/folders', methods=['POST'])
    def create_vault_folder():
        try:
            data = request.get_json(force=True, silent=True) or {}
            name = (data.get('name') or '').strip()
            parent_id = data.get('parent_id')  # None = root folder

            if not name:
                return jsonify({'error': True, 'message': 'Folder name is required.'}), 400
            if len(name) > 80:
                return jsonify({'error': True, 'message': 'Folder name must be under 80 characters.'}), 400

            conn = db
            # Uniqueness check: same name + same parent
            existing = conn.execute(
                'SELECT id FROM vault_folders WHERE name = ? AND (parent_id IS ? OR parent_id = ?)',
                (name, parent_id, parent_id)
            ).fetchone()
            if existing:
                return jsonify({'error': True, 'message': f'A folder named "{name}" already exists here.'}), 409

            c = conn.cursor()
            c.execute(
                'INSERT INTO vault_folders (name, parent_id) VALUES (?, ?)',
                (name, parent_id)
            )
            conn.commit()
            folder_id = c.lastrowid
            return jsonify({'success': True, 'id': folder_id, 'name': name, 'parent_id': parent_id}), 201
        except Exception as e:
            return jsonify({'error': True, 'message': str(e)}), 500

    @app.route('/api/vault/folders/<int:folder_id>', methods=['PATCH', 'OPTIONS'])
    def update_vault_folder(folder_id):
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            data = request.get_json(force=True, silent=True) or {}

            if 'name' in data:
                new_name = (data.get('name') or '').strip()
                if not new_name:
                    return jsonify({'error': True, 'message': 'Folder name cannot be empty.'}), 400
                if len(new_name) > 80:
                    return jsonify({'error': True, 'message': 'Folder name must be under 80 characters.'}), 400
                db.execute('UPDATE vault_folders SET name = ? WHERE id = ?', (new_name, folder_id))
                db.commit()
                return jsonify({'success': True, 'id': folder_id, 'name': new_name}), 200

            if 'parent_id' in data:
                new_parent = data.get('parent_id')  # None = move to root
                # Prevent circular parentage: new_parent must not be self or a descendant
                def get_descendants(fid):
                    kids = [r['id'] for r in db.execute('SELECT id FROM vault_folders WHERE parent_id = ?', (fid,)).fetchall()]
                    result = list(kids)
                    for k in kids:
                        result.extend(get_descendants(k))
                    return result
                if new_parent == folder_id or new_parent in get_descendants(folder_id):
                    return jsonify({'error': True, 'message': 'Cannot move a folder into itself or its own descendant.'}), 400
                db.execute('UPDATE vault_folders SET parent_id = ? WHERE id = ?', (new_parent, folder_id))
                db.commit()
                return jsonify({'success': True, 'id': folder_id, 'parent_id': new_parent}), 200

            return jsonify({'error': True, 'message': 'No valid fields provided (name or parent_id).'}), 400
        except Exception as e:
            return jsonify({'error': True, 'message': str(e)}), 500

    @app.route('/api/vault/folders/<int:folder_id>', methods=['DELETE', 'OPTIONS'])
    def delete_vault_folder(folder_id):
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            def recursive_delete(fid):
                # Delete all documents in this folder
                db.execute('DELETE FROM case_vault WHERE folder_id = ?', (fid,))
                # Get all child folders
                children = [r['id'] for r in db.execute('SELECT id FROM vault_folders WHERE parent_id = ?', (fid,)).fetchall()]
                for child_id in children:
                    recursive_delete(child_id)
                # Delete this folder
                db.execute('DELETE FROM vault_folders WHERE id = ?', (fid,))

            recursive_delete(folder_id)
            db.commit()
            return jsonify({'success': True, 'deleted_id': folder_id}), 200
        except Exception as e:
            return jsonify({'error': True, 'message': str(e)}), 500

    @app.route('/api/vault/documents/<int:doc_id>', methods=['PUT', 'OPTIONS'])
    def update_vault_document(doc_id):
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            data = request.get_json(force=True, silent=True) or {}
            new_content = data.get('content', '')
            db.execute('UPDATE case_vault SET content = ? WHERE id = ?', (new_content, doc_id))
            db.commit()
            return jsonify({'success': True, 'id': doc_id}), 200
        except Exception as e:
            return jsonify({'error': True, 'message': str(e)}), 500

    @app.route('/api/vault/documents/<int:doc_id>', methods=['DELETE', 'OPTIONS'])
    def delete_vault_document(doc_id):
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            db.execute('DELETE FROM case_vault WHERE id = ?', (doc_id,))
            db.commit()
            return jsonify({'success': True, 'deleted_id': doc_id}), 200
        except Exception as e:
            return jsonify({'error': True, 'message': str(e)}), 500

    @app.route('/api/vault/save', methods=['POST', 'OPTIONS'])
    def save_vault_document():
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            data = request.get_json(force=True, silent=True) or {}

            # Virtual Courtroom simulator session save format
            if 'case_facts' in data or 'history' in data:
                case_facts = data.get("case_facts", "")
                history = data.get("history", [])
                title = data.get("title", "VC Session")

                import uuid
                vault_path = os.path.join(os.getcwd(), "vc_vault_data.json")
                if os.path.exists(vault_path):
                    with open(vault_path, "r", encoding="utf-8") as f:
                        vault_data = json.load(f)
                else:
                    vault_data = {}

                case_id = "VC_" + str(uuid.uuid4())[:8]
                vault_data[case_id] = {
                    "title": title,
                    "case_facts": case_facts,
                    "history": history
                }

                with open(vault_path, "w", encoding="utf-8") as f:
                    json.dump(vault_data, f, indent=4)

                return jsonify({"status": "success", "message": f"Session locked to VC Vault (ID: {case_id}).", "id": case_id}), 200

            # Case Vault save format
            case_id     = data.get('case_id')
            title       = data.get('title')
            doc_type    = data.get('doc_type')
            content     = data.get('content')
            folder_id   = data.get('folder_id')    # optional — from SaveToVaultModal
            smart_title = data.get('smart_title')  # optional — auto-generated client-side
            tags_raw    = data.get('tags')          # optional — JSON string or list
            if isinstance(tags_raw, list):
                tags = ','.join(str(t) for t in tags_raw)
            elif isinstance(tags_raw, str):
                try:
                    import json as _json
                    parsed = _json.loads(tags_raw)
                    tags = ','.join(str(t) for t in parsed) if isinstance(parsed, list) else tags_raw
                except Exception:
                    tags = tags_raw
            else:
                tags = None

            # Normalise content — allow empty string (blank draft is valid)
            content = str(content) if content is not None else ''
            if not case_id or not title:
                return jsonify({"error": True, "message": "Missing required fields (case_id, title)."}), 400

            # Format conversion — generate binary blob if requested
            save_format = data.get('format', 'native')
            file_blob = None
            file_format = save_format
            if save_format == 'pdf':
                try:
                    file_blob = generate_pdf_blob(smart_title or title or 'Document', content)
                except Exception as conv_err:
                    return jsonify({'error': True, 'message': f'PDF generation failed: {str(conv_err)}'}), 500
            elif save_format == 'docx':
                try:
                    file_blob = generate_docx_blob(smart_title or title or 'Document', content)
                except Exception as conv_err:
                    return jsonify({'error': True, 'message': f'DOCX generation failed: {str(conv_err)}'}), 500

            audit_messages  = data.get('audit_messages')   # JSON string of [{role, text, ts}]
            session_title_s = data.get('session_title', '')

            # Resolve folder path string for breadcrumb confirmation (best-effort)
            folder_path = ''
            if folder_id:
                try:
                    path_parts = []
                    fid = folder_id
                    seen = set()
                    while fid and fid not in seen:
                        seen.add(fid)
                        row = db.execute('SELECT name, parent_id FROM vault_folders WHERE id = ?', (fid,)).fetchone()
                        if row:
                            path_parts.insert(0, row[0])
                            fid = row[1]
                        else:
                            break
                    folder_path = ' / '.join(path_parts)
                except Exception:
                    pass

            conn = db
            c = conn.cursor()
            c.execute(
                'INSERT INTO case_vault (case_id, title, doc_type, content, folder_id, smart_title, tags, file_blob, file_format) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)',
                (str(case_id), str(title), str(doc_type or ''), str(content), folder_id, smart_title, tags, file_blob, file_format)
            )
            conn.commit()
            inserted_id = c.lastrowid

            # Write AI provenance / audit trail entry
            if audit_messages and inserted_id:
                try:
                    c.execute(
                        'INSERT INTO vault_audit (folder_id, vault_doc_id, session_title, messages_json) VALUES (?, ?, ?, ?)',
                        (folder_id, inserted_id, session_title_s, audit_messages)
                    )
                    conn.commit()
                except Exception:
                    pass

            display_title = smart_title or title
            location_str  = f'{folder_path} / {display_title}' if folder_path else display_title
            return jsonify({
                'success': True,
                'id': inserted_id,
                'message': 'Document successfully saved.',
                'location': location_str,
            }), 200
        except Exception as e:
            return jsonify({"error": True, "message": str(e)}), 500

    @app.route('/api/cases/autocomplete', methods=['GET', 'OPTIONS'])
    def api_cases_autocomplete():
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        q = request.args.get('q', '').strip()
        if not q:
            return jsonify([]), 200
        try:
            conn = db
            c = conn.cursor()
            # Parameterized placeholder — the LIKE wildcard is built into the
            # bound value, never concatenated into the SQL string itself.
            c.execute(
                "SELECT id, title FROM case_vault WHERE title LIKE ? ORDER BY created_at DESC LIMIT 10",
                (f'%{q}%',)
            )
            rows = c.fetchall()
            return jsonify([{"id": r[0], "title": r[1]} for r in rows]), 200
        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({"error": True, "message": str(e)}), 500

    # Citations are stored as a JSON array of {id, title} objects inside
    # case_vault.tags. Both mutations below run as a single atomic SQL
    # UPDATE (no Python-side read-modify-write), wrapped in BEGIN IMMEDIATE
    # so SQLite grabs a RESERVED lock up front instead of deferring — this
    # is what makes concurrent rapid-fire clicks land every citation instead
    # of one write clobbering another. tags may already hold a legacy
    # comma-joined string from /api/vault/save's older format; the
    # json_valid() guard treats anything non-JSON as an empty array rather
    # than erroring or corrupting existing data.
    _CITATION_INSERT_SQL = """
        UPDATE case_vault
        SET tags = (
            SELECT CASE WHEN EXISTS (SELECT 1 FROM json_each(base) WHERE json_extract(value, '$.id') = ?)
                THEN base
                ELSE json_insert(base, '$[#]', json(?))
            END
            FROM (SELECT CASE WHEN tags IS NULL OR tags = '' OR json_valid(tags) = 0 THEN '[]' ELSE tags END AS base
                  FROM case_vault WHERE id = ?)
        )
        WHERE id = ?
    """
    _CITATION_DELETE_SQL = """
        UPDATE case_vault
        SET tags = (
            SELECT json_group_array(json(value))
            FROM json_each((SELECT CASE WHEN tags IS NULL OR tags = '' OR json_valid(tags) = 0 THEN '[]' ELSE tags END
                            FROM case_vault WHERE id = ?))
            WHERE json_extract(value, '$.id') != ?
        )
        WHERE id = ?
    """

    @app.route('/api/vault/documents/<int:doc_id>/citations', methods=['POST', 'OPTIONS'])
    def add_vault_citation(doc_id):
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            data = request.get_json(force=True, silent=True) or {}
            citation_id = str(data.get('id', '')).strip()
            citation_title = str(data.get('title', '')).strip()
            if not citation_id or not citation_title:
                return jsonify({"error": True, "message": "Citation requires both id and title."}), 400

            payload = json.dumps({"id": citation_id, "title": citation_title})

            conn = db
            conn.execute('BEGIN IMMEDIATE')
            row = conn.execute('SELECT id FROM case_vault WHERE id = ?', (doc_id,)).fetchone()
            if not row:
                conn.rollback()
                return jsonify({"error": True, "message": f"Document {doc_id} not found."}), 404
            conn.execute(_CITATION_INSERT_SQL, (citation_id, payload, doc_id, doc_id))
            conn.commit()

            tags_row = conn.execute('SELECT tags FROM case_vault WHERE id = ?', (doc_id,)).fetchone()
            citations = json.loads(tags_row[0]) if tags_row and tags_row[0] else []
            return jsonify({"status": "success", "citations": citations}), 200
        except Exception as e:
            db.rollback()
            import traceback
            traceback.print_exc()
            return jsonify({"error": True, "message": str(e)}), 500

    @app.route('/api/vault/documents/<int:doc_id>/citations/<string:citation_id>', methods=['DELETE', 'OPTIONS'])
    def remove_vault_citation(doc_id, citation_id):
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            conn = db
            conn.execute('BEGIN IMMEDIATE')
            row = conn.execute('SELECT id FROM case_vault WHERE id = ?', (doc_id,)).fetchone()
            if not row:
                conn.rollback()
                return jsonify({"error": True, "message": f"Document {doc_id} not found."}), 404
            conn.execute(_CITATION_DELETE_SQL, (doc_id, citation_id, doc_id))
            conn.commit()

            tags_row = conn.execute('SELECT tags FROM case_vault WHERE id = ?', (doc_id,)).fetchone()
            citations = json.loads(tags_row[0]) if tags_row and tags_row[0] else []
            return jsonify({"status": "success", "citations": citations}), 200
        except Exception as e:
            db.rollback()
            import traceback
            traceback.print_exc()
            return jsonify({"error": True, "message": str(e)}), 500

    @app.route('/api/vault/documents/<int:doc_id>/download', methods=['GET', 'OPTIONS'])
    def download_vault_document(doc_id):
        from flask import Response as FlaskResponse
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            row = db.execute(
                'SELECT file_blob, file_format, smart_title, title FROM case_vault WHERE id = ?', (doc_id,)
            ).fetchone()
            if not row or not row[0]:
                return jsonify({'error': True, 'message': 'No binary file stored for this document.'}), 404
            fmt = row[1] or 'native'
            name = row[2] or row[3] or f'document_{doc_id}'
            if fmt == 'pdf':
                mime = 'application/pdf'
                ext  = '.pdf'
            else:
                mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                ext  = '.docx'
            safe_name = ''.join(c for c in name if c.isalnum() or c in ' _-.')
            return FlaskResponse(
                bytes(row[0]),
                mimetype=mime,
                headers={'Content-Disposition': f'attachment; filename="{safe_name}{ext}"'},
            )
        except Exception as e:
            return jsonify({'error': True, 'message': str(e)}), 500

    @app.route('/api/vault/audit-trail', methods=['GET', 'OPTIONS'])
    def vault_audit_trail():
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            folder_id = request.args.get('folder_id', type=int)
            conn = db
            old_rf = conn.row_factory
            conn.row_factory = sqlite3.Row
            try:
                c = conn.cursor()
                if folder_id is not None:
                    c.execute('''
                        SELECT va.id, va.folder_id, va.vault_doc_id, va.session_title,
                               va.messages_json, va.created_at,
                               cv.title AS doc_title, cv.doc_type,
                               vf.name  AS folder_name
                        FROM   vault_audit va
                        LEFT JOIN case_vault    cv ON va.vault_doc_id = cv.id
                        LEFT JOIN vault_folders vf ON va.folder_id    = vf.id
                        WHERE  va.folder_id = ?
                        ORDER  BY va.created_at DESC
                    ''', (folder_id,))
                else:
                    c.execute('''
                        SELECT va.id, va.folder_id, va.vault_doc_id, va.session_title,
                               va.messages_json, va.created_at,
                               cv.title AS doc_title, cv.doc_type,
                               vf.name  AS folder_name
                        FROM   vault_audit va
                        LEFT JOIN case_vault    cv ON va.vault_doc_id = cv.id
                        LEFT JOIN vault_folders vf ON va.folder_id    = vf.id
                        ORDER  BY va.created_at DESC
                        LIMIT  60
                    ''')
                rows = c.fetchall()
                threads = []
                for row in rows:
                    d = dict(row)
                    try:
                        d['messages'] = json.loads(d.get('messages_json') or '[]')
                    except Exception:
                        d['messages'] = []
                    threads.append(d)
                return jsonify({'threads': threads}), 200
            finally:
                conn.row_factory = old_rf
        except Exception as e:
            return jsonify({'error': True, 'message': str(e)}), 500

    @app.route('/api/vault/documents', methods=['GET'])
    def get_vault_documents():
        try:
            q = (request.args.get('q') or '').strip()
            folder_id_param = request.args.get('folder_id')
            try:
                limit = max(1, min(int(request.args.get('limit', 30)), 100))
            except (TypeError, ValueError):
                limit = 30
            try:
                offset = max(0, int(request.args.get('offset', 0)))
            except (TypeError, ValueError):
                offset = 0

            where_clauses = []
            params = []
            if q:
                like = f'%{q}%'
                where_clauses.append('(cv.title LIKE ? OR cv.content LIKE ? OR cv.case_id LIKE ?)')
                params.extend([like, like, like])
            elif folder_id_param is not None:
                if folder_id_param in ('root', 'null', ''):
                    where_clauses.append('cv.folder_id IS NULL')
                else:
                    try:
                        where_clauses.append('cv.folder_id = ?')
                        params.append(int(folder_id_param))
                    except ValueError:
                        pass
            where_sql = ('WHERE ' + ' AND '.join(where_clauses)) if where_clauses else ''

            conn = db
            old_row_factory = conn.row_factory
            conn.row_factory = sqlite3.Row
            try:
                c = conn.cursor()
                total = c.execute(
                    f"SELECT COUNT(*) FROM case_vault cv {where_sql}", params
                ).fetchone()[0]

                # CRITICAL MEMORY SAFEGUARD: LIMIT/OFFSET is applied in SQL so
                # only one bounded page of rows is ever pulled out of SQLite —
                # `content` can be hundreds of KB per judgment, and this table
                # has 900+ rows, so a plain fetchall() of everything (as the
                # old query did) loads the entire vault's text into Python
                # memory on every request regardless of what's on screen.
                c.execute(
                    f"""
                    SELECT cv.*,
                           vf.name       AS folder_name,
                           vf.parent_id  AS folder_parent_id
                    FROM   case_vault cv
                    LEFT JOIN vault_folders vf ON cv.folder_id = vf.id
                    {where_sql}
                    ORDER BY cv.created_at DESC
                    LIMIT ? OFFSET ?
                    """,
                    params + [limit, offset]
                )
                rows = c.fetchall()
                dict_rows = []
                for row in rows:
                    d = dict(row)
                    d['title'] = resolve_vault_title(d.get('title'), d.get('case_id'), d.get('content'))
                    dict_rows.append(d)
            finally:
                conn.row_factory = old_row_factory
            return jsonify({"documents": dict_rows, "total": total, "limit": limit, "offset": offset}), 200
        except Exception as e:
            return jsonify({"error": True, "message": str(e)}), 500

    @app.route('/api/vault/meta', methods=['GET', 'OPTIONS'])
    def get_vault_meta():
        """Lightweight document metadata for the Save-to-Vault modal explorer.
        Returns id, title, doc_type, folder_id, file_format, created_at — NO content field."""
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            conn = db
            old_rf = conn.row_factory
            conn.row_factory = sqlite3.Row
            try:
                rows = conn.execute(
                    'SELECT id, title, doc_type, folder_id, file_format, created_at '
                    'FROM case_vault ORDER BY created_at DESC'
                ).fetchall()
                docs = [dict(r) for r in rows]
            finally:
                conn.row_factory = old_rf
            return jsonify({'documents': docs}), 200
        except Exception as e:
            return jsonify({'error': True, 'message': str(e)}), 500

    @app.route('/api/firm-library', methods=['GET', 'OPTIONS'])
    def get_firm_library():
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            conn = db
            old_rf = conn.row_factory
            conn.row_factory = sqlite3.Row
            try:
                # Schema inspection rather than a hardcoded column list — if
                # case_vault is ever migrated to rename/drop a column, this
                # route degrades to the case_id fallback below instead of
                # 500ing on "no such column".
                table_cols = {row[1] for row in conn.execute("PRAGMA table_info('case_vault')").fetchall()}
                descriptive_col = next(
                    (c for c in ('case_name', 'description', 'title', 'smart_title') if c in table_cols),
                    None
                )

                select_cols = ['id', 'case_id', 'created_at']
                if descriptive_col and descriptive_col not in select_cols:
                    select_cols.append(descriptive_col)
                for optional_col in ('validity_status', 'ratio_headnote', 'content'):
                    if optional_col in table_cols:
                        select_cols.append(optional_col)
                query = f"SELECT {', '.join(select_cols)} FROM case_vault ORDER BY created_at DESC"
                rows = conn.execute(query).fetchall()

                docs = []
                for r in rows:
                    raw_title = r[descriptive_col] if descriptive_col else None
                    timestamp = r['created_at']
                    row_keys = r.keys()
                    docs.append({
                        'id': r['id'],
                        'case_id': r['case_id'],
                        'title': format_vault_title(raw_title, r['case_id']),
                        'timestamp': timestamp,
                        'last_updated': timestamp,
                        'updated': timestamp and str(timestamp).split(' ')[0] or '',
                        'type': 'Precedent',
                        'category': 'Precedent',
                        'author': 'Internal Vault',
                        'validity_status': (r['validity_status'] if 'validity_status' in row_keys else None) or 'Green',
                        'ratio_headnote': (r['ratio_headnote'] if 'ratio_headnote' in row_keys else None) or '',
                        # Truncated — this is a list endpoint returning every
                        # entry at once; the accordion reader only needs a
                        # preview, not the full document (a large judgment's
                        # full text here would bloat every single list
                        # fetch, not just the one row someone expands).
                        'content': ((r['content'] if 'content' in row_keys else None) or '')[:4000],
                    })
            finally:
                conn.row_factory = old_rf
            return jsonify(docs), 200
        except Exception as e:
            return jsonify({'error': True, 'message': str(e)}), 500

    @app.route('/api/firm-library', methods=['POST'])
    def create_firm_library_entry():
        """Legal Forms Library -> Firm Library persistence ("Save to Library
        & Exit"). Writes the drafted TipTap HTML into case_vault — the same
        table GET /api/firm-library already reads from — so a saved draft
        shows up in the Firm Library list on the very next fetch, no
        separate storage concept required."""
        try:
            data = request.get_json(silent=True) or {}
            title = (data.get('title') or '').strip()
            html = data.get('html') or ''
            category = (data.get('category') or 'Draft').strip()

            if not title:
                return jsonify({'error': True, 'message': 'Title is required.'}), 400
            if not html.strip():
                return jsonify({'error': True, 'message': 'No document content provided.'}), 400

            case_id = f"draft_{int(time.time())}_{uuid.uuid4().hex[:6]}"
            conn = db
            c = conn.cursor()
            c.execute(
                'INSERT INTO case_vault (case_id, title, doc_type, content) VALUES (?, ?, ?, ?)',
                (case_id, title, f'Firm Library Draft — {category}', html),
            )
            conn.commit()
            inserted_id = c.lastrowid

            # Best-effort local index write — never blocks/fails the save
            # itself (index_document already swallows its own errors), so
            # a slow-to-load embedding model can't turn a document save
            # into a 500.
            try:
                from utils.local_search import index_document
                index_document(inserted_id, title, html)
            except Exception as e:
                print(f"[create_firm_library_entry] indexing skipped: {e}")

            row = conn.execute(
                'SELECT id, case_id, title, created_at FROM case_vault WHERE id = ?', (inserted_id,)
            ).fetchone()
            timestamp = row[3] if row else None
            return jsonify({
                'id': inserted_id,
                'case_id': case_id,
                'title': title,
                'timestamp': timestamp,
                'last_updated': timestamp,
                'updated': timestamp and str(timestamp).split(' ')[0] or '',
                'type': category,
                'category': category,
                'author': 'Internal Vault',
                'validity_status': 'Green',
                'ratio_headnote': '',
            }), 201
        except Exception as e:
            return jsonify({'error': True, 'message': str(e)}), 500

    @app.route('/api/firm-library/external-search', methods=['GET', 'OPTIONS'])
    def firm_library_external_search():
        """LLM-free semantic lookup over the Pinecone-backed case-law index,
        for the Firm Library's 'External Database' toggle. Deliberately
        never touches Groq — this is a fast browse/search grid, not the
        AI-synthesized Dual-Brain research dossier (/api/legal-research)."""
        if request.method == 'OPTIONS':
            return jsonify({}), 200

        query = (request.args.get('query') or '').strip()
        if not query:
            return jsonify({'error': True, 'message': 'Query cannot be empty.'}), 400

        try:
            pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
            index = _get_pinecone_index(pc)
            # Hardcoded, not read from PINECONE_NAMESPACE — this route only
            # ever targets one namespace, and reading it from an env var
            # that could be unset-to-empty-string in some deployment would
            # silently query Pinecone's separate default namespace instead
            # (empty string is itself a valid, DIFFERENT namespace, not
            # "unset") and return zero results with no error.
            namespace = "legal-cases"

            embed_response = pc.inference.embed(
                model=EMBED_MODEL,
                inputs=[query],
                parameters={"input_type": "query", "truncate": "END"},
            )
            query_vector = embed_response[0].values

            # top_k deliberately high (45, not the 10 we actually return) —
            # a single case can contribute dozens of chunks, so a small
            # top_k would rarely survive dedup as more than one or two
            # distinct cases; over-fetching lets the loop below collect 10
            # actually-DIFFERENT cases instead of 10 near-duplicate chunks
            # of the same one or two. No score-threshold filtering — every
            # match is a dedup/formatting candidate regardless of score.
            pinecone_results = index.query(
                vector=query_vector,
                top_k=45,
                include_metadata=True,
                namespace=namespace,
            )

            results = []
            seen_cases = set()
            for match in (pinecone_results.matches or []):
                meta = match.metadata or {}
                case_id = meta.get('case_id') or match.id
                # CRITICAL DEDUPLICATION — same as /api/legal-research: a
                # case can have hundreds of chunks, so without this the
                # first case in the ranking would flood the grid with N
                # identical-looking rows.
                if case_id in seen_cases:
                    continue
                seen_cases.add(case_id)

                # worker.py's actual ingested metadata is only ever
                # {case_id, year, chunk_index, text} — title/case_name/
                # citation/etc. below are NOT present today. The extra
                # fallback keys are deliberate future-proofing (bulletproof
                # against the metadata schema gaining fields later), not
                # evidence they exist now; the real, working fallback is
                # format_case_title(case_id), not a flat placeholder string
                # that would otherwise render identically for every row.
                resolved_title = (
                    meta.get('title') or meta.get('case_name') or meta.get('doc_title')
                    or format_case_title(case_id)
                )
                text = meta.get('snippet') or meta.get('text') or meta.get('chunk_text') or ''
                year = meta.get('year') or meta.get('date') or ''

                results.append({
                    # case_id, not match.id (a raw per-CHUNK vector hash) —
                    # this is what /api/document/<id> and Pin-to-Vault
                    # expect, and what the dedup above is keyed on; using
                    # the chunk id here would silently break both.
                    'id': case_id,
                    'case_id': case_id,
                    'title': resolved_title,
                    'case_title': resolved_title,
                    'case_name': resolved_title,
                    'snippet': text[:200],
                    'text': text[:200],
                    'citation': meta.get('citation') or meta.get('reporter') or 'Supreme Court of India',
                    'year': str(year) if year else 'N/A',
                    'score': round(float(match.score), 4) if match.score is not None else None,
                    # Fallback values so this ALSO still maps onto the
                    # {id, title, category, updated, author} schema the
                    # Firm Library table's shared rows already render for
                    # internal entries — no real category/author exists
                    # for a raw Pinecone chunk.
                    'category': 'Case Law',
                    'updated': f"{year}-01-01" if year else '',
                    'author': 'External Case Law DB',
                    'tags': [],
                })
                if len(results) >= 10:
                    break

            return jsonify({'status': 'success', 'results': results, 'data': results}), 200
        except Exception as e:
            print(f"[PINECONE ERROR] {e}")
            return jsonify({"status": "error", "message": "Vector database unavailable.", "results": []}), 200

    # "I'm Feeling Lucky" proxy for the Firm Library's "Open in Indian
    # Kanoon" link — a plain <a href> click can't carry an Authorization
    # header, so this is deliberately unauthenticated, same as the other
    # Firm Library routes above. The redirect target is always a fixed
    # indiankanoon.org path built from either a regex-extracted digit-only
    # doc id or urllib.parse.quote(query) — query never controls the
    # destination HOST, only the search term, so this isn't an open redirect.
    # Kanoon links a result to /doc/<id>/ normally, but to
    # /docfragment/<id>/?formInput=... whenever the search term was matched
    # inside the document body rather than the title — both point at the
    # exact same underlying document id, so both forms must match here or
    # every keyword-in-body hit is silently dropped.
    _KANOON_DOC_ID_RE = re.compile(r'/doc(?:fragment)?/(\d+)')
    _KANOON_CASE_SPLIT_RE = re.compile(r'\s+(?:v\.|vs\.?|versus)\s+', re.IGNORECASE)
    _KANOON_WORD_RE = re.compile(r'[A-Za-z0-9]+')
    # Em dash, en dash, hyphen-minus, curly double/single quotes, straight
    # quotes, and non-breaking space — Kanoon's search silently returns
    # zero results for a formInput containing any of these (routine in a
    # pasted heading/citation, e.g. "Indian Contract Act — Free Consent
    # Doctrine"). Not a raw string — \xa0 needs Python to actually decode
    # it into the real non-breaking-space character; a raw string would
    # leave it as the four literal characters backslash-x-a-0 instead. The
    # plain "-" is placed last in the class so it's never misread as a
    # range operator against a neighbouring \uXXXX-style escape.
    _KANOON_DASH_QUOTE_RE = re.compile('[—–“”‘’"\'\xa0-]')
    _KANOON_BRACKETS_RE = re.compile(r'\([^)]*\)|\[[^\]]*\]')
    _KANOON_YEAR_PAREN_RE = re.compile(r'\(\s*(\d{4})\s*\)')
    _KANOON_REPORTER_RE = re.compile(r'\b(SCC|AIR|SCR|INSC)\b', re.IGNORECASE)
    # The exact list the task specifies for party-token extraction, reused
    # everywhere a stopword filter is needed (party tokens, cleaning a
    # non-case query, tokenizing hit text for scoring) rather than
    # inventing a second, undocumented list for those other uses.
    _KANOON_STOPWORDS = {
        'state', 'union', 'india', 'ltd', 'pvt', 'anr', 'ors', 'the', 'of', 'co', 'corp',
    }
    # Matches the START of the citation apparatus that routinely trails a
    # pasted party name — "State of Kerala (1973) 4 SCC 225" — so it can be
    # truncated off before tokenizing. Without this, "SCC"/"225" get treated
    # as party-name tokens and pollute the title:() clause with citation
    # noise no judgment title ever actually contains, defeating the whole
    # point of title-scoping.
    _KANOON_CITATION_TAIL_RE = re.compile(r'\(\s*\d{4}\s*\)|\b\d{4}\b|\b(?:SCC|AIR|SCR|INSC)\b', re.IGNORECASE)

    def _kanoon_strip_citation_tail(text):
        match = _KANOON_CITATION_TAIL_RE.search(text or '')
        return text[:match.start()] if match else (text or '')

    def _kanoon_strip_brackets(text):
        # Drops "(BALCO)"/"[Regn. No. 4]"-style asides entirely, including
        # their contents — these are abbreviations/annotations, not part of
        # the actual party name or search text.
        return _KANOON_BRACKETS_RE.sub(' ', text or '')

    def _kanoon_tokenize(text):
        return [
            tok for tok in _KANOON_WORD_RE.findall(text or '')
            if tok.lower() not in _KANOON_STOPWORDS and len(tok) > 1
        ]

    def _kanoon_extract_boosters(raw_query):
        """Citation signals pulled from the query BEFORE any bracket
        cleanup — the year lives inside the parens _kanoon_strip_brackets
        would otherwise delete. Safe to call with the dash/quote-cleaned
        query (parens are untouched by that cleanup), just not one that's
        already had its brackets stripped. Stored case-neutral/uppercase so
        a single upper-cased substring check at scoring time covers both
        booster types."""
        boosters = set()
        for year in _KANOON_YEAR_PAREN_RE.findall(raw_query):
            boosters.add(f"({year})")
        for tag in _KANOON_REPORTER_RE.findall(raw_query):
            boosters.add(tag.upper())
        return boosters

    def _kanoon_title_clause(tokens):
        if not tokens:
            return ''
        return f"title:({' AND '.join(tokens)})"

    def _kanoon_fetch_results(form_input):
        """One Kanoon search request -> its parsed result_title elements,
        or None on any network/HTTP failure. None must always be treated by
        the caller as 'give up, redirect to fallback_url' — never retried
        silently, per the network guardrail.

        Routed through ZenRows (premium proxy, 'in' exit node) whenever
        ZENROWS_API_KEY is configured — production runs from a
        datacenter/cloud-host IP that Kanoon routinely blocks or degrades,
        which no amount of header-spoofing from that host gets around.
        Falls back to a direct request only when the key is absent, i.e.
        local development from an ordinary residential IP.
        """
        zenrows_key = os.getenv("ZENROWS_API_KEY")
        target_url = f"https://indiankanoon.org/search/?formInput={urllib.parse.quote(form_input)}"

        try:
            if zenrows_key:
                resp = requests.get(
                    "https://api.zenrows.com/v1/",
                    # ZenRows' own param name is "apikey" (no underscore) —
                    # confirmed live: sending "api_key" gets a 401 AUTH001
                    # ("No apikey provided") on every single request, which
                    # would silently degrade every production redirect to
                    # the bare search page instead of a specific doc.
                    params={
                        'apikey': zenrows_key,
                        'url': target_url,
                        'premium_proxy': 'true',
                        'proxy_country': 'in',
                    },
                    timeout=15,
                )
            else:
                resp = requests.get(
                    "https://indiankanoon.org/search/",
                    params={'formInput': form_input},
                    timeout=5,
                    headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                                            'Chrome/122.0.0.0 Safari/537.36'},
                )
            # A 403/5xx response body is still valid HTML (an error/block
            # page) — requests won't raise on its own, so raise_for_status()
            # is what actually routes a block or outage into the
            # RequestException branch below instead of silently parsing
            # zero results out of a block page.
            resp.raise_for_status()
        except requests.exceptions.RequestException as err:
            print(f"[KANOON ERROR] Proxy Mode ({'ZenRows' if zenrows_key else 'Direct'}): {err}")
            return None

        # Matched by class only, not class+tag — Kanoon's search page was
        # redesigned since this was first written; the result title element
        # is an <h4> today (it used to be a <div>), and tying the query to
        # one specific tag would silently match zero results the next time
        # the markup shifts again.
        return BeautifulSoup(resp.text, 'html.parser').find_all(class_='result_title')

    @app.route('/api/kanoon-redirect', methods=['GET'])
    def kanoon_redirect():
        query = request.args.get('query', '') or ''
        # fallback_url intentionally stays on the raw, unaltered query —
        # it's the last-resort "go search it yourself" link, so it should
        # reflect exactly what the user typed, not our cleaned-up version.
        fallback_url = f"https://indiankanoon.org/search/?formInput={urllib.parse.quote(query)}"

        # Kanoon's search silently returns zero results for a formInput
        # containing unicode dashes/smart-quotes/NBSP — routine in a pasted
        # heading or citation ("Indian Contract Act — Free Consent
        # Doctrine"). Cleaned BEFORE is_case detection: an NBSP standing in
        # for a normal space around "v."/"vs." could otherwise dodge the
        # \s+ boundaries in the case-split regex.
        cleaned_query = _KANOON_DASH_QUOTE_RE.sub(' ', query)
        cleaned_query = re.sub(r'\s+', ' ', cleaned_query).strip()

        is_case = bool(_KANOON_CASE_SPLIT_RE.search(cleaned_query))

        if not is_case:
            # Statutory/doctrine queries ("Doctrine of Promissory Estoppel")
            # almost never repeat their exact phrasing inside any single
            # judgment's title+snippet, so the token-intersection scorer
            # below routinely scores every real, relevant hit as 0 and
            # punts to the bare search page even when Kanoon's own #1
            # result is a perfectly good answer — confirmed empirically:
            # a batch of realistic doctrine queries scored 0 across the
            # board despite Kanoon returning 10 real results each. Skip
            # scoring entirely here and trust Kanoon's own ranking instead.
            try:
                result_titles = _kanoon_fetch_results(cleaned_query)
                if result_titles:
                    a_tag = result_titles[0].find('a', href=True)
                    if a_tag:
                        # /doc(?:fragment)?/ — not a bare /doc/ — matches both
                        # of Kanoon's link forms: plain /doc/<id>/ when the
                        # match was in the title, /docfragment/<id>/ when it
                        # was in the body. A /doc/-only pattern silently drops
                        # to the fallback for the /docfragment/ case — and that
                        # case is the COMMON one here: sampled against several
                        # realistic doctrine queries, 3 of 4 top results linked
                        # via /docfragment/, not /doc/.
                        doc_match = _KANOON_DOC_ID_RE.search(a_tag['href'])
                        if doc_match:
                            return redirect(f"https://indiankanoon.org/doc/{doc_match.group(1)}/")
            except Exception as e:
                # A Kanoon markup change or unexpected parse error here must
                # still land the user on a working search page, not Flask's
                # default HTML 500 — this route backs a plain <a href> link
                # click, which can't render a JSON error to the user anyway.
                print(f"[KANOON REDIRECT ERROR] non-case branch: {e}")
            return redirect(fallback_url)

        citation_boosters = _kanoon_extract_boosters(cleaned_query)
        party_a_raw, party_b_raw = _KANOON_CASE_SPLIT_RE.split(cleaned_query, maxsplit=1)
        tokens_a = _kanoon_tokenize(_kanoon_strip_brackets(_kanoon_strip_citation_tail(party_a_raw)))[:3]
        tokens_b = _kanoon_tokenize(_kanoon_strip_brackets(_kanoon_strip_citation_tail(party_b_raw)))[:3]
        query_tokens = {t.lower() for t in tokens_a + tokens_b}
        clauses = [c for c in (_kanoon_title_clause(tokens_a), _kanoon_title_clause(tokens_b)) if c]
        # Degenerate case: stopword-filtering emptied BOTH parties (e.g.
        # "The State v. The Union") — title:() is invalid Lucene syntax,
        # so fall back to the cleaned query instead of sending Kanoon
        # malformed input.
        form_input = ' '.join(clauses) if clauses else cleaned_query

        # Attempt 1: title-scoped search.
        result_titles = _kanoon_fetch_results(form_input)
        if result_titles is None:
            return redirect(fallback_url)

        # Attempt 2: title:() scoping can legitimately return zero hits for a
        # case whose title doesn't literally contain both party names as
        # formatted (abbreviations, "& Ors" variants, etc.) — one plain retry
        # with the cleaned (not raw) query before giving up; retrying with
        # the UNCLEANED original would just reproduce the same
        # unicode-choking failure the cleanup step exists to fix. Skipped
        # when form_input already IS the cleaned query (a case query
        # degenerate enough that both parties fell back to it) — retrying
        # would just re-fetch the identical 0 results.
        if not result_titles and form_input != cleaned_query:
            result_titles = _kanoon_fetch_results(cleaned_query)
            if result_titles is None:
                return redirect(fallback_url)

        best_doc_id = None
        best_score = 0

        try:
            for title_el in result_titles:
                link = title_el.find('a', href=True)
                if not link:
                    continue
                doc_match = _KANOON_DOC_ID_RE.search(link['href'])
                if not doc_match:
                    continue
                doc_id = doc_match.group(1)

                # separator=' ' is required: Kanoon bolds individual matched
                # words ("Constitution of <b>India</b>"), and get_text()'s
                # default empty separator glues adjacent fragments across
                # tag boundaries into "ofIndia" — silently breaking any
                # token that crosses a bold-tag boundary.
                title_text = title_el.get_text(separator=' ', strip=True)

                # The snippet lives in a sibling <div class="headline"> next
                # to the <h4 class="result_title">, both inside a shared
                # <article class="result"> wrapper — fall back to searching
                # the whole result container in case a future markup shift
                # reorders the siblings.
                snippet_el = title_el.find_next_sibling(class_='headline')
                if snippet_el is None:
                    result_container = title_el.find_parent(class_='result')
                    snippet_el = result_container.find(class_='headline') if result_container else None
                snippet_text = snippet_el.get_text(separator=' ', strip=True) if snippet_el else ''
                combined_text = f"{title_text} {snippet_text}"

                hit_tokens = {t.lower() for t in _KANOON_WORD_RE.findall(combined_text)}
                score = len(query_tokens & hit_tokens)
                if score == 0:
                    continue

                # Citation multiplier: collapse whitespace immediately
                # inside parens first — a bolded citation year renders as
                # "(<b>1973</b>)", and get_text's separator=' ' would
                # otherwise insert stray spaces ("( 1973 )") that break a
                # literal "(1973)" substring check.
                normalized = re.sub(r'\(\s+', '(', combined_text)
                normalized = re.sub(r'\s+\)', ')', normalized).upper()
                if any(booster in normalized for booster in citation_boosters):
                    score *= 3

                if score > best_score:
                    best_score = score
                    best_doc_id = doc_id
        except Exception:
            return redirect(fallback_url)

        if best_doc_id and best_score > 0:
            return redirect(f"https://indiankanoon.org/doc/{best_doc_id}/")

        return redirect(fallback_url)

    # Register blueprints
    from routes.auth_routes import auth_bp
    from routes.sso_routes import sso_bp
    from routes.library_routes import library_bp
    from routes.ai_routes import ai_bp
    from routes.document_routes import doc_bp
    from routes.contract_routes import contract_bp
    from routes.argument_routes import argument_bp
    from routes.court_routes import court_bp
    from routes.causelist_routes import causelist_bp
    from routes.client_routes import client_bp
    from routes.billing_routes import billing_bp
    from routes.conflict_routes import conflict_bp
    from routes.team_routes import team_bp

    app.register_blueprint(court_bp)
    app.register_blueprint(argument_bp)
    app.register_blueprint(causelist_bp)
    app.register_blueprint(client_bp)
    app.register_blueprint(billing_bp)
    app.register_blueprint(conflict_bp)
    app.register_blueprint(team_bp)
    app.register_blueprint(auth_bp, url_prefix='/api/auth')
    app.register_blueprint(sso_bp)  # url_prefix already baked into sso_bp's own definition
    app.register_blueprint(library_bp)  # url_prefix already baked into library_bp's own definition
    app.register_blueprint(ai_bp, url_prefix='/api/ai')
    app.register_blueprint(doc_bp, url_prefix='/api/documents')
    app.register_blueprint(contract_bp, url_prefix='/api/contract')

    # Keep-alive health check — always 200 while the server is up.
    # The client clears its interval on 401/403 from other protected endpoints.
    @app.route('/api/ping')
    def api_ping():
        return jsonify({'ok': True})

    def force_pristine_groq_messages(messages):
        clean = []
        allowed_keys = {'role', 'content', 'name', 'tool_calls', 'tool_call_id'}
        for m in messages:
            if not isinstance(m, dict):
                continue
            # Explicitly extract only safe API parameters
            sanitized_msg = {k: v for k, v in m.items() if k in allowed_keys}
            
            # Handle potential nested dictionary mutations inside tool_calls if they exist
            if 'tool_calls' in sanitized_msg and sanitized_msg['tool_calls'] is None:
                del sanitized_msg['tool_calls']
                
            clean.append(sanitized_msg)
        return clean

    @app.route('/api/legal-research', methods=['POST', 'OPTIONS'])
    def api_legal_research():
        if request.method == 'OPTIONS':
            return jsonify({}), 200

        data = request.get_json(force=True, silent=True) or {}
        lawyer_question = data.get('question', '').strip()

        if not lawyer_question:
            return jsonify({"error": True, "message": "Question cannot be empty."}), 400

        try:
            # 1. Initialize Pinecone & Groq using existing .env keys
            pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
            index = _get_pinecone_index(pc)
            namespace = os.getenv("PINECONE_NAMESPACE", "legal-cases")
            groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

            # 2. Embed the question and query the raw vector index — worker.py
            # ingests via index.upsert() with real metadata (case_id, year,
            # chunk_index, text), not via Pinecone's integrated-inference
            # records API, so matches carry .metadata directly instead of
            # needing the id string parsed apart.
            embed_response = pc.inference.embed(
                model=EMBED_MODEL,
                inputs=[lawyer_question],
                parameters={"input_type": "query", "truncate": "END"},
            )
            query_vector = embed_response[0].values

            # top_k=10 (not 2) deliberately over-fetches chunks so dedup below
            # has multiple same-case chunks to actually collapse — a small
            # top_k would rarely return more than one chunk per case anyway,
            # making the dedup step a no-op.
            pinecone_results = index.query(
                vector=query_vector,
                top_k=10,
                include_metadata=True,
                namespace=namespace,
            )

            context_chunks = []
            sources = []
            seen_cases = set()

            for match in (pinecone_results.matches or []):
                metadata = match.metadata or {}
                case_id = metadata.get('case_id') or match.id

                text = metadata.get('text', '')
                if text:
                    context_chunks.append(text)

                # CRITICAL DEDUPLICATION — a case can have hundreds of chunks;
                # without this, the same case would flood the UI as N
                # identical-looking source cards.
                if case_id in seen_cases:
                    continue
                seen_cases.add(case_id)

                sources.append({
                    "case_id": case_id,
                    "title": format_case_title(case_id),
                    "year": metadata.get('year', ''),
                    "snippet": text[:200],
                })

            context_text = "\n\n".join(context_chunks)

            if not context_chunks:
                return jsonify({
                    "status": "success",
                    "answer": "I couldn't find any relevant precedents in the Firm Library regarding this query.",
                    "sources": []
                }), 200

            # 3. Generate Answer using Groq
            system_prompt = (
                "You are an expert senior legal counsel in India. Analyze the following legal fragments "
                "and answer the user's question precisely based *only* on the context provided. If the "
                "context doesn't contain the answer, state that clearly."
            )

            user_prompt = f"Context:\n{context_text}\n\nQuestion: {lawyer_question}"

            chat_completion = groq_client.chat.completions.create(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                model="llama-3.1-8b-instant",
                temperature=0.2,
            )

            ai_answer = chat_completion.choices[0].message.content

            return jsonify({
                "status": "success",
                "answer": ai_answer,
                "sources": sources
            }), 200

        except Exception as e:
            print(f"[PINECONE ERROR] {e}")
            return jsonify({
                "answer": "I am currently unable to access the legal database. Please try again later.",
                "context": []
            }), 200

    @app.route('/api/legal-research/document/<string:case_id>', methods=['GET', 'OPTIONS'])
    def api_legal_research_document(case_id):
        if request.method == 'OPTIONS':
            return jsonify({}), 200

        # Strict allow-list validation — case_id arrives as a raw URL path
        # segment, so this is the only thing standing between an attacker and
        # a crafted "../" or SQL-metacharacter payload reaching the query.
        if not CASE_ID_RE.match(case_id):
            return jsonify({"error": True, "message": "Invalid case_id format."}), 400

        try:
            conn = db
            c = conn.cursor()
            c.execute(
                "SELECT title, content FROM case_vault WHERE case_id = ? ORDER BY created_at DESC LIMIT 1",
                (case_id,)
            )
            row = c.fetchone()

            if not row:
                return jsonify({"error": True, "message": f"Document with case_id '{case_id}' not found."}), 404

            title, content = row
            return jsonify({"title": title, "full_text": content}), 200

        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({"error": True, "message": str(e)}), 500

    @app.route('/api/document/<string:case_id>', methods=['GET', 'OPTIONS'])
    def api_document_reconstruct(case_id):
        """Reconstructs a full worker.py-ingested judgment from its Pinecone
        chunks. Distinct from /api/legal-research/document/<case_id> above,
        which reads AI-drafted documents out of the local case_vault table —
        S3-ingested judgments were never written there, only into Pinecone.
        """
        if request.method == 'OPTIONS':
            return jsonify({}), 200

        if not PINECONE_CASE_ID_RE.match(case_id):
            return jsonify({"error": True, "message": "Invalid case_id format."}), 400

        try:
            pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
            index = pc.Index(host=os.getenv("PINECONE_HOST"))
            namespace = os.getenv("PINECONE_NAMESPACE", "legal-cases")

            resolved_case_id = case_id
            if MD5_HASH_RE.match(case_id):
                # The caller only has a raw chunk vector id (e.g. a search
                # result whose metadata lacked case_id) — resolve the real
                # case_id from that one vector's metadata first.
                fetch_result = index.fetch(ids=[case_id], namespace=namespace)
                vector = fetch_result.vectors.get(case_id)
                if vector and vector.metadata:
                    resolved_case_id = vector.metadata.get('case_id', case_id)

            # Dummy near-zero vector: this is a filter-only lookup, not a
            # similarity search — the score is discarded, so the vector's
            # direction doesn't matter, but a true [0.0]*N is undefined
            # under this index's cosine metric (division by zero) and would
            # error instead of just returning arbitrarily-ranked matches.
            query_result = index.query(
                vector=DUMMY_QUERY_VECTOR,
                filter={"case_id": resolved_case_id},
                top_k=3000,  # a large judgment can run to hundreds of chunks
                include_metadata=True,
                namespace=namespace,
            )

            matches = query_result.matches or []
            if not matches:
                return jsonify({"error": True, "message": f"Document with case_id '{resolved_case_id}' not found."}), 404

            # int() cast is required: Pinecone metadata typing isn't
            # guaranteed uniform across every historical ingestion run, and
            # a string sort would order chunks as 1, 10, 11, 2, 3... instead
            # of 1, 2, 3... 10, 11, silently scrambling the reassembled text.
            sorted_matches = sorted(
                matches,
                key=lambda m: int((m.metadata or {}).get('chunk_index', 0))
            )
            content = "\n\n".join((m.metadata or {}).get('text', '') for m in sorted_matches)

            return jsonify({
                "case_id": resolved_case_id,
                "title": format_case_title(resolved_case_id),
                "content": content,
            }), 200

        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({"error": True, "message": f"Failed to reconstruct document: {str(e)}"}), 500

    @app.route('/api/ai/summarize-judgment', methods=['POST', 'OPTIONS'])
    def api_summarize_judgment():
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            data = request.get_json(force=True, silent=True) or {}
            text = str(data.get('text', '')).strip()
            if not text:
                return jsonify({"error": True, "message": "text cannot be empty."}), 400

            # Legal preamble blindspot: the first ~6000 characters of a judgment
            # are almost always procedural history and counsel appearances, not
            # the court's actual reasoning. Sample the opening (for structural
            # context) plus the closing (where the ratio decidendi and ultimate
            # holding live) instead of naively truncating from the top. Guard
            # against the two windows overlapping on short judgments, which
            # would otherwise duplicate the opening into the context block.
            if len(text) <= 8000:
                context_block = text
            else:
                context_block = f"{text[:2000]}\n\n...[middle of judgment omitted]...\n\n{text[-6000:]}"

            groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))
            system_prompt = (
                "You are an expert senior legal counsel in India. Summarize the following judgment "
                "excerpt into exactly 3 concise bulleted legal takeaways. Focus on the court's ratio "
                "decidendi and ultimate holding — ignore procedural history and counsel appearances "
                "even if they dominate the excerpt."
            )
            chat_completion = groq_client.chat.completions.create(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": context_block},
                ],
                model="llama-3.1-8b-instant",
                temperature=0.2,
            )
            summary = chat_completion.choices[0].message.content
            return jsonify({"status": "success", "summary": summary}), 200
        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({"error": True, "message": str(e)}), 500

    @app.route('/api/chat', methods=['POST', 'OPTIONS'])
    def api_chat():
        if request.method == 'OPTIONS':
            return jsonify({}), 200
        try:
            # 1. Safe Payload Parsing with force=True to handle missing/wrong Content-Type header
            data = request.get_json(force=True, silent=True) or {}
            query = data.get('query', '').strip()
            current_path = data.get('currentPath', '/')
            params = data.get('params', {})

            # Backward compatibility fallback
            if not query:
                query = data.get('question', '').strip()

            # 2. Dynamic Payload & DoS Protection (max 1500 chars)
            if len(query) > 1500:
                return jsonify({"error": True, "message": "Query exceeds maximum limit of 1500 characters."}), 400

            if not query:
                return jsonify({"error": True, "message": "No query or question provided."}), 400

            # Securely load GROQ_API_KEY from environment
            groq_api_key = os.getenv("GROQ_API_KEY")
            if not groq_api_key:
                raise ValueError("GROQ_API_KEY environment variable is not set.")
            
            # Initialize Groq Client
            client = Groq(api_key=groq_api_key)

            # 3. Anchored System Prompt & Context Injection
            system_instruction = (
                "You are LexAI, an AI Legal Associate and junior counsel embedded inside LexAmplify — an intelligent legal practice management platform for advocates practising under Indian law.\n\n"

                "YOUR PERSONA:\n"
                "You are a highly capable junior advocate trained in Indian law: IPC, CrPC, CPC, Constitution of India, IBC, Companies Act, Negotiable Instruments Act, Transfer of Property Act, IP law (Trade Marks Act, Patents Act, Copyright Act), Consumer Protection Act, Arbitration & Conciliation Act, and procedural law across all tiers of Indian courts.\n"
                "You think and respond like a diligent junior lawyer: precise, professionally cautious, well-researched, and always deferential to the senior advocate (the user). "
                "You draft documents in correct legal English, cite relevant provisions, identify procedural requirements, flag risks, and manage deadlines — all under the advocate's supervision.\n\n"

                "DAILY TASKS YOU HANDLE (examples — not exhaustive):\n"
                "- Draft petitions, plaints, written statements, bail applications, legal notices, NDAs, MOUs, agreements, affidavits, vakalatnamas\n"
                "- Research IPC/CrPC/CPC sections, Supreme Court and High Court precedents, bare act provisions\n"
                "- Navigate to any part of LexAmplify: High Courts, District Courts, Contract Analyzer, Conflict Engine, Calendar, Case Vault, War Room\n"
                "- Schedule hearings, drop-dead deadlines, tickler reminders, and court appearances in the Legal Calendar\n"
                "- Analyze uploaded contracts for risk clauses\n"
                "- Prepare hearing briefs and arguments for courtroom simulation\n"
                "- Save drafted documents to the Case Vault for the advocate's records\n\n"

                "NAVIGATION CAPABILITY:\n"
                "You can navigate to any feature of LexAmplify. Valid routes include:\n"
                "  /dashboard, /contract-analyzer, /court-resources, /conflict-engine, /calendar, /vault, /war-room\n"
                "When navigating to /court-resources, you may also specify a tab: supreme, highcourt, district, laws, forms, events, courtfee, enotary, iptracker.\n"
                "Use the navigate tool when the user gives a navigation command.\n\n"

                "CRITICAL EXECUTION MANDATE (TOOL CALLING):\n"
                "1. Use native JSON tool calling — never output raw XML tags.\n"
                "2. Execute EXACTLY ONE tool per turn. Never chain tools.\n"
                "3. Always provide ALL required parameters. Never pass empty {} objects.\n\n"

                "WORKFLOW MANDATE (THE 3-PHASE PIPELINE):\n"
                "You orchestrate a strict state machine: Draft → Approval → Simulate.\n\n"

                "PHASE 1 — DRAFTING:\n"
                "When the advocate requests a document, invoke the document generation tool, present the draft, and STOP. Never assume approval.\n\n"

                "PHASE 2 — HUMAN-IN-THE-LOOP APPROVAL:\n"
                "You cannot save to the Case Vault yourself. Wait for the advocate to approve or reject via the UI.\n"
                "  If Rejected: acknowledge and ask for revised parameters.\n"
                "  If Approved: confirm the document is now saved in the Case Vault.\n\n"

                "PHASE 3 — SIMULATION:\n"
                "Never launch a courtroom simulation from conversational memory alone. Always retrieve the approved document from the Case Vault first via the appropriate tool.\n\n"

                "ANTI-CHAINING GUARDRAIL:\n"
                "For multi-step commands like 'Draft, save, and simulate', execute only Phase 1. "
                "Explicitly tell the advocate you are pausing for their UI approval before proceeding further."
            )

            # Define Tool Schema for Groq SDK
            tools = [
                {
                    "type": "function",
                    "function": {
                        "name": "propose_calendar_events",
                        "description": (
                            "Propose deadlines and tickler warning events (e.g. 30, 14, or 7 days prior to a critical date) "
                            "based on user scheduling requests. STRICT RULE: If you have no events to propose, do not invoke this tool."
                        ),
                        "parameters": {
                            "type": "object",
                            "properties": {
                                "events": {
                                    "type": "array",
                                    "description": "The list of proposed events.",
                                    "items": {
                                        "type": "object",
                                        "properties": {
                                            "event_date": {
                                                "type": "string",
                                                "description": "The date of the event in YYYY-MM-DD format (e.g. '2026-06-01')."
                                            },
                                            "event_type": {
                                                "type": "string",
                                                "enum": ["drop_dead", "tickler", "appearance", "task"],
                                                "description": "The type of the event."
                                            },
                                            "title": {
                                                "type": "string",
                                                "description": "The title or description of the event."
                                            },
                                            "related_case_id": {
                                                "type": "string",
                                                "description": "The related case ID if applicable, or an empty string."
                                            }
                                        },
                                        "required": ["event_date", "event_type", "title", "related_case_id"]
                                    }
                                }
                            }
                        }
                    }
                },
                {
                    "type": "function",
                    "function": {
                        "name": "navigate_ui",
                        "description": "Navigate the user interface to a specific page or route in the application.",
                        "parameters": {
                            "type": "object",
                            "properties": {
                                "target_route": {
                                    "type": "string",
                                    "description": "The target route to navigate to (e.g. '/dashboard')."
                                }
                            },
                            "required": ["target_route"]
                        }
                    }
                },
                {
                    "type": "function",
                    "function": {
                        "name": "propose_document_draft",
                        "description": "Propose a draft of a legal document (e.g. bail application, petition, agreement) for user review and approval.",
                        "parameters": {
                            "type": "object",
                            "properties": {
                                "case_id": {
                                    "type": "string",
                                    "description": "The related case ID."
                                },
                                "title": {
                                    "type": "string",
                                    "description": "The title of the document draft."
                                },
                                "doc_type": {
                                    "type": "string",
                                    "description": "The document type (e.g. 'Bail Application', 'Contract')."
                                },
                                "content": {
                                    "type": "string",
                                    "description": "The complete text content of the drafted legal document."
                                }
                            },
                            "required": ["case_id", "title", "doc_type", "content"]
                        }
                    }
                },
                {
                    "type": "function",
                    "function": {
                        "name": "search_case_vault",
                        "description": "Search the Case Vault. If you know the exact Case ID, provide it in case_id. If you do not know the exact ID, provide a detailed query in search_query (e.g., \"bail application for chain snatching\"). DO NOT guess or hallucinate case_id numbers.",
                        "parameters": {
                            "type": "object",
                            "properties": {
                                "case_id": {
                                    "type": ["string", "null"],
                                    "description": "The specific case identifier. If the user does not specify a case, you may pass null. If you pass null, you MUST ask the user 'Which case vault would you like me to search?'"
                                },
                                "search_query": {
                                    "type": ["string", "null"],
                                    "description": "A semantic text query to find the document by keywords if the case_id is unknown."
                                }
                            }
                        }
                    }
                },
                {
                    "type": "function",
                    "function": {
                        "name": "run_courtroom_simulation",
                        "description": "Run a 5-stage agentic courtroom simulation based on a saved document in the Case Vault.",
                        "parameters": {
                            "type": "object",
                            "properties": {
                                "case_id": {
                                    "type": "string",
                                    "description": "The case ID of the document in the Case Vault."
                                },
                                "client_side": {
                                    "type": "string",
                                    "description": "The client side, e.g., 'Appellant', 'Respondent', 'Plaintiff', 'Defendant'."
                                }
                            },
                            "required": ["case_id", "client_side"]
                        }
                    }
                }
            ]

            incoming_messages = data.get('messages', [])
            
            if incoming_messages:
                clean_messages = force_pristine_groq_messages(incoming_messages)
            else:
                clean_messages = [{"role": "user", "content": f"User query: {query}"}]
                
            messages = [{"role": "system", "content": system_instruction}] + clean_messages

            # 4. The First Groq Call (Decision Phase)
            sanitized_messages = force_pristine_groq_messages(messages)
            
            response = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=sanitized_messages,
                tools=tools,
                tool_choice="auto"
            )

            # 5. Interceptor Logic (Context Preservation)
            response_message = response.choices[0].message
            tool_calls = response_message.tool_calls

            if tool_calls:
                tool_call = tool_calls[0]
                tool_name = tool_call.function.name
                
                # Dynamic Argument Parsing Guardrail
                try:
                    args = json.loads(tool_call.function.arguments)
                    # Dynamically check required fields against our active schema
                    active_schema = next((t["function"] for t in tools if t["function"]["name"] == tool_name), None)
                    if active_schema and "parameters" in active_schema:
                        required_fields = active_schema["parameters"].get("required", [])
                        missing_fields = [f for f in required_fields if f not in args]
                        if missing_fields:
                            raise ValueError(f"Missing required parameters: {missing_fields}")
                except Exception as e:
                    # Intercept before hitting external gateways and feed error back to LLM
                    error_feedback = (
                        f"System Guardrail Alert: Your tool invocation for '{tool_name}' failed validation. "
                        f"Error: {str(e)}. "
                        "You MUST retry your function call and ensure you provide all required parameters exactly as specified in the schema."
                    )
                    
                    try:
                        assistant_msg = response_message.model_dump()
                    except AttributeError:
                        try:
                            assistant_msg = response_message.dict()
                        except AttributeError:
                            assistant_msg = {
                                "role": "assistant",
                                "content": response_message.content,
                                "tool_calls": [{"id": tool_call.id, "type": "function", "function": {"name": tool_call.function.name, "arguments": tool_call.function.arguments}}]
                            }
                    
                    messages.append(assistant_msg)
                    messages.append({
                        "role": "tool",
                        "tool_call_id": tool_call.id,
                        "name": tool_name,
                        "content": error_feedback
                    })
                    
                    retry_sanitized = force_pristine_groq_messages(messages)
                    retry_response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=retry_sanitized
                    )
                    
                    fallback_text = retry_response.choices[0].message.content or "Tool validation failed. Please try again."
                    return jsonify({
                        "action": "chat",
                        "message": fallback_text,
                        "answer": fallback_text
                    }), 200

                if tool_name == "run_courtroom_simulation":
                    case_id = args.get('case_id')
                    client_side = args.get('client_side', 'Appellant')
                    
                    conn = db
                    c = conn.cursor()
                    c.execute("SELECT content FROM case_vault WHERE case_id = ? ORDER BY created_at DESC LIMIT 1", (str(case_id),))
                    row = c.fetchone()
                    if not row or not row[0]:
                        return jsonify({"error": True, "message": f"Document with Case ID '{case_id}' not found in the Case Vault."}), 400
                        
                    content = row[0]
                    truncated_content = content[:8000]
                    
                    stage1_prompt = (
                        "Analyze the following legal document content:\n\n"
                        f"{truncated_content}\n\n"
                        "Task:\n"
                        "1. Extract the top 3 core legal issues in this case.\n"
                        "2. Summarize these issues into a single 5-word search query string.\n\n"
                        "You must respond with a JSON object. The JSON object must contain exactly these keys:\n"
                        '- "extracted_issues": "A summary string listing the top 3 core legal issues."\n'
                        '- "search_query": "A single 5-word search query string."'
                    )
                    
                    res_stage1 = client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=force_pristine_groq_messages([{"role": "user", "content": stage1_prompt}]),
                        response_format={"type": "json_object"}
                    )
                    
                    stage1_data = extract_json_from_llm_response(res_stage1.choices[0].message.content)
                    if not isinstance(stage1_data, dict):
                        return jsonify({
                            "error": True,
                            "message": "Stage 1 parsing failed: AI returned an unparseable response.",
                            "code": "LLM_JSON_PARSE_ERROR",
                        }), 502
                        
                    stage1_text = stage1_data.get("extracted_issues", "")
                    search_query = stage1_data.get("search_query", "")
                    
                    tavily_results = []
                    tavily_key = os.getenv("TAVILY_API_KEY")
                    if tavily_key:
                        try:
                            tavily_client = TavilyClient(api_key=tavily_key)
                            query = f"Indian Supreme Court landmark judgments regarding {search_query} site:indiankanoon.org"
                            search_result = tavily_client.search(query=query, search_depth="advanced", max_results=3)
                            results = search_result.get("results", [])
                            for r in results:
                                tavily_results.append({
                                    "title": r.get('title', ''),
                                    "snippet": r.get('content', ''),
                                    "url": r.get('url', '')
                                })
                        except Exception as e:
                            print(f"Tavily search failed: {e}")
                            tavily_results = []
                            
                    stage3_prompt = (
                        f"Act as an Indian Advocate representing the {client_side}.\n"
                        "Draft a structured opening argument for your case.\n"
                        "You MUST cite the provided web search cases to back up your claims if they are relevant.\n\n"
                        f"Facts and Issues:\n{stage1_text}\n\n"
                        f"Live Web Search Cases Retrieved:\n{json.dumps(tavily_results, indent=2)}\n\n"
                        "Provide only the drafted opening argument text. Do not include conversational text or markdown code blocks."
                    )
                    
                    res_stage3 = client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=force_pristine_groq_messages([{"role": "user", "content": stage3_prompt}])
                    )
                    stage3_text = res_stage3.choices[0].message.content or ""
                    
                    stage4_prompt = (
                        "Act as the opposing counsel in this litigation.\n"
                        f"Here is the opening argument presented by the {client_side}:\n\n"
                        f"{stage3_text}\n\n"
                        "Identify legal weaknesses and generate aggressive counter-questions. For each question, provide a suggested rebuttal for the other side to defend themselves.\n\n"
                        "You must respond with a JSON object. The JSON object must contain exactly one key:\n"
                        '- "opposing_counter_questions": A list of objects, where each object contains exactly:\n'
                        '    - "question": "The opposing counsel\'s objection or counter-question string."\n'
                        '    - "suggested_rebuttal": "A suggested defense rebuttal string."'
                    )
                    
                    res_stage4 = client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=force_pristine_groq_messages([{"role": "user", "content": stage4_prompt}]),
                        response_format={"type": "json_object"}
                    )
                    
                    stage4_json = extract_json_from_llm_response(res_stage4.choices[0].message.content)
                    if not isinstance(stage4_json, dict):
                        # Stages 1-3 already succeeded and are real, usable
                        # data — only the red-team sub-section failed to
                        # parse, so the overall response still degrades
                        # gracefully (with the failure visibly marked in the
                        # fallback question itself) rather than discarding
                        # everything the pipeline already produced.
                        stage4_json = {"opposing_counter_questions": [{"question": "Failed to parse opponent arguments.", "suggested_rebuttal": res_stage4.choices[0].message.content}]}

                    return jsonify({
                        "action": "simulate_courtroom",
                        "simulationData": {
                            "client_side": client_side,
                            "extracted_issues": stage1_text,
                            "live_citations": tavily_results,
                            "opening_argument": stage3_text,
                            "red_team": stage4_json
                        },
                        "message": "Dynamic simulation complete. Transferring to the War Room..."
                    }), 200

                elif tool_name == "propose_calendar_events":
                    proposed_events = args.get('events', [])
                    return jsonify({
                        "action": "confirm_schedule",
                        "proposed_events": proposed_events,
                        "message": "I have drafted the schedule and ticklers. Please review and approve.",
                        "answer": "I have drafted the schedule and ticklers. Please review and approve."
                    }), 200

                elif tool_name == "propose_document_draft":
                    # INTERCEPT IMMEDIATELY. Do NOT write to DB.
                    return jsonify({
                        "action": "review_document",
                        "draft": args,
                        "message": "Draft ready. Please review and approve to save to the Vault.",
                        "answer": "Draft ready. Please review and approve to save to the Vault."
                    }), 200

                elif tool_name == "search_case_vault":
                    conn = db
                    old_row_factory = conn.row_factory
                    conn.row_factory = sqlite3.Row
                    try:
                        c = conn.cursor()
                        case_id = args.get("case_id")
                        search_query = args.get("search_query")
                        
                        if case_id and str(case_id).lower() != "null":
                            c.execute("SELECT * FROM case_vault WHERE case_id = ? ORDER BY created_at DESC", (str(case_id),))
                            rows = c.fetchall()
                        elif search_query:
                            c.execute("SELECT * FROM case_vault ORDER BY created_at DESC")
                            all_docs = c.fetchall()
                            stop_words = {"a", "an", "the", "and", "or", "but", "if", "for", "with", "about", "as", "by", "in", "to", "of", "on", "is", "are"}
                            keywords = [w.lower() for w in search_query.split() if w.lower() not in stop_words]
                            
                            best_doc = None
                            max_overlap = 0
                            
                            for doc in all_docs:
                                doc_text = (str(doc['title']) + " " + str(doc['content'])).lower()
                                overlap = sum(1 for kw in keywords if kw in doc_text)
                                if overlap > max_overlap:
                                    max_overlap = overlap
                                    best_doc = doc
                            
                            rows = [best_doc] if best_doc else []
                        else:
                            c.execute("SELECT * FROM case_vault ORDER BY created_at DESC")
                            rows = c.fetchall()
                    finally:
                        conn.row_factory = old_row_factory
                    
                    results = []
                    for row in rows:
                        results.append(
                            f"Document ID: {row['id']}\n"
                            f"Case ID: {row['case_id']}\n"
                            f"Title: {row['title']}\n"
                            f"Type: {row['doc_type']}\n"
                            f"Created At: {row['created_at']}\n"
                            f"Content:\n{row['content']}\n"
                            "---"
                        )
                    
                    if results:
                        formatted_vault_results = "\n".join(results)
                    else:
                        if args.get("search_query"):
                            formatted_vault_results = "No documents found matching this query in the Case Vault."
                        else:
                            formatted_vault_results = "No documents found in the Case Vault."

                    # Convert response_message to dict to be safe or use model_dump
                    try:
                        assistant_msg = response_message.model_dump()
                    except AttributeError:
                        try:
                            assistant_msg = response_message.dict()
                        except AttributeError:
                            assistant_msg = {
                                "role": "assistant",
                                "content": response_message.content,
                                "tool_calls": [
                                    {
                                        "id": tool_call.id,
                                        "type": "function",
                                        "function": {
                                            "name": tool_call.function.name,
                                            "arguments": tool_call.function.arguments
                                        }
                                    }
                                ]
                            }

                    messages.append(assistant_msg)
                    messages.append({
                        "role": "tool",
                        "tool_call_id": tool_call.id,
                        "name": "search_case_vault",
                        "content": formatted_vault_results
                    })

                    sanitized_second_messages = force_pristine_groq_messages(messages)

                    second_response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=sanitized_second_messages
                    )
                    final_text = second_response.choices[0].message.content or ""
                    return jsonify({
                        "action": "chat",
                        "message": final_text,
                        "answer": final_text
                    }), 200

                elif tool_name == "navigate_ui":
                    # INTERCEPT IMMEDIATELY. Do NOT make a second Groq call.
                    target_route = args.get("target_route")
                    
                    # Titanium Route Validation
                    valid_routes = ['/dashboard', '/contract-analyzer', '/court-resources', '/conflict-engine', '/calendar', '/vault']
                    if target_route in valid_routes:
                        return jsonify({
                            "action": "navigate",
                            "target_route": target_route,
                            "message": "Navigating to requested module...",
                            "answer": "Navigating to requested module..."
                        }), 200
                    else:
                        return jsonify({
                            "action": "chat",
                            "message": "I cannot navigate to that specific module.",
                            "answer": "I cannot navigate to that specific module."
                        }), 200
                else:
                    # Graceful fallback for any other tool/hallucination
                    return jsonify({
                        "action": "chat",
                        "message": "I cannot perform that action right now.",
                        "answer": "I cannot perform that action right now."
                    }), 200
            else:
                final_text = response_message.content or ""
                # 4. Fallback: Standard chat (no tools called)
                return jsonify({
                    "action": "chat",
                    "message": final_text,
                    "answer": final_text
                }), 200

        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({"error": True, "message": str(e)}), 500

    @app.route('/api/calendar/save', methods=['POST'])
    def save_calendar_events():
        try:
            data = request.get_json(force=True, silent=True) or {}
            events = data.get('events', [])
            
            conn = db
            c = conn.cursor()
            for event in events:
                event_date = event.get('event_date')
                event_type = event.get('event_type')
                title = event.get('title')
                related_case_id = event.get('related_case_id')
                location = event.get('location', '')
                opposing_counsel = event.get('opposing_counsel', '')
                c.execute('''
                    INSERT INTO calendar_events (event_date, event_type, title, related_case_id, location, opposing_counsel)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (event_date, event_type, title, related_case_id, location, opposing_counsel))
            
            conn.commit()
            return jsonify({"success": True, "message": "Events successfully saved."}), 200
        except Exception as e:
            return jsonify({"error": True, "message": str(e)}), 500

    @app.route('/api/calendar/events', methods=['GET'])
    def get_calendar_events_all():
        try:
            conn = db
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            c.execute("SELECT * FROM calendar_events ORDER BY event_date ASC")
            rows = c.fetchall()
            dict_rows = [dict(row) for row in rows]
            return jsonify({"events": dict_rows}), 200
        except Exception as e:
            return jsonify({"error": True, "message": str(e)}), 500

    @app.route('/api/ecourts/sync', methods=['POST'])
    def sync_ecourts_cnr():
        try:
            data = request.get_json(force=True, silent=True) or {}
            cnr = (data.get('cnr') or '').strip()
            # eCourts API integration point — simulated response until live API key is provisioned
            return jsonify({
                "success": True,
                "cnr": cnr,
                "hearings_added": 3,
                "message": f"Matter Synced: 3 Hearings added to Calendar."
            }), 200
        except Exception as e:
            return jsonify({"error": True, "message": str(e)}), 500

    @app.route('/api/districts', methods=['GET'])
    def get_districts():
        try:
            base_dir = os.path.dirname(os.path.abspath(__file__))
            primary_path = os.path.join(base_dir, 'data', 'districts.json')
            backup_path  = os.path.join(base_dir, 'data', 'districts_backup.json')
            target_file  = primary_path if os.path.exists(primary_path) else backup_path
            with open(target_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return jsonify(data)
        except FileNotFoundError:
            return jsonify({"error": "Database not found. Please run scraper.py first."}), 404

    # Page routes
    @app.route('/')
    def landing():
        return render_template('landing.html')

    @app.route('/login')
    def login_page():
        return render_template('login.html')

    @app.route('/signup')
    def signup_page():
        return render_template('signup.html')

    @app.route('/dashboard')
    def dashboard():
        return render_template('dashboard.html')

    @app.route('/contract-viewer')
    def contract_viewer():
        return render_template('contract_viewer.html')

    with app.app_context():
        # Explicitly import all models to ensure SQL schemas are registered in metadata
        from models.user import User
        from models.case import Case
        from models.document import Document
        sqlalchemy_db.create_all()

    # --- START DATABASE BUILDER ---
    import sqlite3
    try:
        # Connect to your exact database file
        conn = sqlite3.connect('lex_assistant.db')
        c = conn.cursor()
        
        # Build the Case Vault table
        c.execute('''
            CREATE TABLE IF NOT EXISTS case_vault (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                case_id TEXT,
                title TEXT,
                doc_type TEXT,
                content TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # Build the Calendar Events table
        c.execute('''
            CREATE TABLE IF NOT EXISTS calendar_events (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                event_date TEXT,
                event_type TEXT,
                title TEXT,
                related_case_id TEXT,
                location TEXT,
                opposing_counsel TEXT
            )
        ''')
        # Migrate existing DB if new columns are missing
        for _col in ('location', 'opposing_counsel'):
            try:
                conn.execute(f'ALTER TABLE calendar_events ADD COLUMN {_col} TEXT')
            except Exception:
                pass

        # Vault folders table + case_vault migration
        conn.execute('''
            CREATE TABLE IF NOT EXISTS vault_folders (
                id        INTEGER PRIMARY KEY AUTOINCREMENT,
                name      TEXT NOT NULL,
                parent_id INTEGER REFERENCES vault_folders(id) ON DELETE CASCADE,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        for _col, _type in (
            ('folder_id', 'INTEGER'), ('smart_title', 'TEXT'), ('tags', 'TEXT'),
            ('file_blob', 'BLOB'), ('file_format', 'TEXT'),
            # Firm Library dense-table columns — SQLite ADD COLUMN can't take
            # a non-constant default via ALTER for existing rows in older
            # sqlite builds, but a literal string default IS allowed and
            # backfills every existing row to 'Green' (matches
            # verify_citation_status()'s baseline assumption: unverified
            # case law is presumed good law until a scrape says otherwise).
            ('validity_status', "TEXT DEFAULT 'Green'"), ('ratio_headnote', 'TEXT'),
        ):
            try:
                conn.execute(f'ALTER TABLE case_vault ADD COLUMN {_col} {_type}')
            except Exception:
                pass
        conn.execute('''
            CREATE TABLE IF NOT EXISTS vault_audit (
                id             INTEGER PRIMARY KEY AUTOINCREMENT,
                folder_id      INTEGER REFERENCES vault_folders(id) ON DELETE SET NULL,
                vault_doc_id   INTEGER,
                session_title  TEXT,
                messages_json  TEXT,
                created_at     DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')

        # Build the Document Chunks table
        c.execute('''
            CREATE TABLE IF NOT EXISTS document_chunks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                case_id INTEGER,
                document_id INTEGER NOT NULL,
                chunk_index INTEGER NOT NULL,
                chunk_text TEXT NOT NULL,
                embedding TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        conn.commit()
        conn.close()
        print("Database tables verified successfully!")
    except Exception as e:
        print(f"Database setup error: {e}")
    # --- END DATABASE BUILDER ---
    return app

if __name__ == '__main__':
    init_db()
    app = create_app()
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)
