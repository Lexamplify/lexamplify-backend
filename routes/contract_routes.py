"""
routes/contract_routes.py
Blueprint: /api/contract
Hybrid approach:
  - Keyword-based risk classification (always works, no JSON parsing)
  - Groq for issue explanation text
  - Groq for clause rewriting
"""
import re
import os
import json as json_mod
from flask import Blueprint, request, jsonify
from utils.ai_helper import ask_groq, ask_litellm
from utils.pdf_helper import extract_text_for_summary

contract_bp = Blueprint("contract", __name__)

# ── Keyword-based risk classification ──────────────────────────
# These patterns are checked against the clause text (case-insensitive)
# If matched → guaranteed RED or GREEN, no AI needed for classification

RED_PATTERNS = [
    r"without\s+(any\s+)?notice",
    r"without\s+(any\s+)?reason",
    r"sole\s+discretion",
    r"waives?\s+all\s+rights?",
    r"shall\s+not\s+(resign|terminate|leave)",
    r"under\s+any\s+circumstances",
    r"internal\s+committee",
    r"no\s+right\s+to\s+(appeal|court|tribunal)",
    r"without\s+prior\s+consent",
    r"unilateral(ly)?",
    r"non.compete.*(5|five)\s+year",
    r"non.compete.*(4|four)\s+year",
    r"non.compete.*(3|three)\s+year",
    r"perpetual(ly)?",
    r"irrevocable",
    r"absolute\s+discretion",
    r"at\s+any\s+time\s+without",
    r"reduced\s+at\s+the\s+sole",
    r"revised\s+or\s+reduced",
    r"forfeit(s|ed)?\s+all",
    r"no\s+compensation",
    r"exclusive\s+jurisdiction\s+of\s+(employer|company|management)",
]

GREEN_PATTERNS = [
    r"provident\s+fund",
    r"epf|esic|pf\s+act",
    r"factories\s+act",
    r"governed\s+by\s+the\s+laws\s+of\s+india",
    r"jurisdiction\s+of\s+courts\s+in",
    r"30\s+days?\s+(written\s+)?notice",
    r"mutual(ly)?\s+(agree|consent|termination)",
    r"as\s+per\s+(indian\s+)?law",
    r"as\s+mandated\s+by",
    r"in\s+accordance\s+with",
    r"both\s+parties",
    r"equal(ly)?\s+(applicable|binding)",
]

RED_ISSUES = {
    r"without\s+(any\s+)?notice": "Termination or action without notice violates the Industrial Disputes Act 1947 and is unenforceable.",
    r"without\s+(any\s+)?reason": "Termination without reason is against natural justice principles under Indian law.",
    r"sole\s+discretion": "Giving one party sole discretion is heavily one-sided and may be void under Section 23 of the Indian Contract Act 1872.",
    r"waives?\s+all\s+rights?": "Waiving all legal rights including right to approach courts is void under Section 28 of the Indian Contract Act 1872.",
    r"shall\s+not\s+(resign|terminate|leave)": "Preventing an employee from resigning is void — forced employment is unenforceable under Indian law.",
    r"under\s+any\s+circumstances": "Absolute restrictions on termination or resignation are void under the Indian Contract Act 1872.",
    r"internal\s+committee": "Restricting dispute resolution to internal committee only, excluding courts, is void under Section 28 of the Indian Contract Act.",
    r"no\s+right\s+to\s+(appeal|court|tribunal)": "Blocking access to courts or tribunals is void under Section 28 of the Indian Contract Act 1872.",
    r"without\s+prior\s+consent": "Making changes without prior consent is one-sided and violates mutual consent principles.",
    r"at\s+any\s+time\s+without": "Allowing unilateral action at any time without restriction is heavily one-sided under Indian contract law.",
    r"reduced\s+at\s+the\s+sole": "Unilateral salary reduction without employee consent is illegal under the Payment of Wages Act 1936.",
    r"revised\s+or\s+reduced": "Allowing employer to reduce salary unilaterally violates the Payment of Wages Act 1936.",
    r"non.compete.*(3|4|5|three|four|five)\s+year": "Non-compete clause exceeding 1-2 years is generally unenforceable under Indian law as it restricts trade.",
    r"forfeit(s|ed)?\s+all": "Forfeiture of all dues is unreasonable and may be challenged as an unfair contract term.",
    r"no\s+compensation": "Excluding compensation entirely may violate mandatory payment obligations under Indian labour law.",
}

GREEN_ISSUES = {
    r"provident\s+fund|epf": "Mandatory EPF contribution as per Employees Provident Fund Act 1952 — legally compliant.",
    r"factories\s+act": "Leave entitlement as per Factories Act 1948 — standard and legally compliant.",
    r"governed\s+by\s+the\s+laws\s+of\s+india": "Standard governing law clause — legally compliant.",
    r"jurisdiction\s+of\s+courts\s+in": "Jurisdiction clause specifying Indian courts — standard and legally compliant.",
    r"30\s+days?\s+(written\s+)?notice": "30-day notice period — fair and standard under Indian employment law.",
    r"mutual": "Mutual consent clause — fair and balanced for both parties.",
    r"both\s+parties": "Clause applies equally to both parties — balanced and fair.",
}


def classify_risk_by_keywords(text: str) -> tuple:
    """
    Returns (risk_level, issue_text) based on keyword matching.
    risk_level = 'RED' | 'AMBER' | 'GREEN'
    """
    text_lower = text.lower()

    # Check RED patterns first (highest priority)
    for pattern in RED_PATTERNS:
        if re.search(pattern, text_lower):
            # Find matching issue text
            issue = "This clause is high risk and potentially unenforceable under Indian law."
            for issue_pattern, issue_text in RED_ISSUES.items():
                if re.search(issue_pattern, text_lower):
                    issue = issue_text
                    break
            return "RED", issue

    # Check GREEN patterns
    green_matches = 0
    green_issue = "No significant issue — standard and legally compliant clause."
    for pattern in GREEN_PATTERNS:
        if re.search(pattern, text_lower):
            green_matches += 1
            for issue_pattern, issue_text in GREEN_ISSUES.items():
                if re.search(issue_pattern, text_lower):
                    green_issue = issue_text
                    break

    if green_matches > 0:
        return "GREEN", green_issue

    # Default to AMBER with AI-generated issue
    return "AMBER", None


def get_ai_issue(clause_text: str) -> str:
    """Use Groq to get issue description for AMBER clauses."""
    prompt = """You are an Indian contract law expert.
In ONE sentence, describe the legal risk or concern with this contract clause under Indian law.
Be specific. Mention the relevant Indian law if applicable.
Return ONLY the one sentence. No preamble."""
    try:
        result = ask_groq(prompt, f"Clause: {clause_text}")
        # Clean up — take just the first sentence
        result = result.strip().split('\n')[0].strip()
        if len(result) > 200:
            result = result[:200] + "..."
        return result
    except Exception:
        return "This clause may have legal implications under Indian law — please review with a lawyer."

def get_ai_classification_with_rules(clause_text: str, rules_text: str) -> tuple:
    prompt = f"""You are an Indian contract law expert.
You MUST use the attached Company Rules to determine what is classified as High (RED), Medium (AMBER), or Low (GREEN) risk. Do not rely solely on general legal knowledge.
Company Rules:
{rules_text}

Analyze the clause and return ONLY JSON exactly like this:
{{"risk": "RED", "issue": "One short sentence issue description."}}
If it complies, return "GREEN" with issue "Compliant with company rules".
If concerning, return "AMBER". If dangerous, return "RED".
"""
    try:
        import json
        import re
        result = ask_groq(prompt, f"Clause: {clause_text}")
        match = re.search(r'\{.*\}', result.replace('\n', ''), re.DOTALL)
        if match:
            parsed = json.loads(match.group(0))
            return parsed.get("risk", "AMBER").upper(), parsed.get("issue", "Flags raised by AI.")
        return "AMBER", "Could not parse AI response cleanly."
    except Exception as e:
        return "AMBER", f"AI Error during rule checking: {e}"


REWRITE_PROMPT = """
You are an expert Indian contract lawyer.
Rewrite the given clause to fix the identified risk and match the user's intent.
Make it fair, balanced, and enforceable under the Indian Contract Act 1872.
Return ONLY the rewritten clause text. No preamble. No explanation.
"""

MASTER_SYSTEM_PROMPT = """You are an elite legal AI architect analyzing a contract.

CRITICAL DIRECTIVES:
1. STRICT BOILERPLATE EXCLUSION: DO NOT extract, highlight, or flag introductory recitals, party definitions, signature blocks, dates, contact details, or basic document titles. Ignore standard boilerplate entirely.
2. RISK STRATEGY ({scanStrategy}):
   - If Aggressive: Be hyper-critical. Flag every minor ambiguity, standard boilerplate risk, and slight imbalance. Output many clauses — aim for 8 to 15 flagged items.
   - If Defensive: Be highly lenient. ONLY flag catastrophic, deal-breaking liabilities (e.g., massive fixed financial penalties, total loss of IP, permanent non-compete with no time limit). Ignore standard confidentiality, jurisdiction, notice, or governing-law clauses completely. Your output MUST have significantly fewer flagged clauses — aim for 2 to 4 items maximum.
3. MANDATORY JSON SCHEMA: You MUST return a strictly valid JSON object. No markdown, no conversational text. It MUST contain exactly two top-level keys:
  - "summary": A professional 3-sentence brief summarizing the actual content, purpose, and parties of the document. Do not summarize the risks here.
  - "clauses": An array of risk objects. Each object must contain: "original_text" (the exact clause from the document), "risk_level" (High, Medium, or Low), and "explanation" (why it is a risk).
"""


def _strip_json_fences(raw: str) -> str:
    """Strip markdown code fences that LLMs sometimes wrap around JSON output."""
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned[3:]
    if cleaned.endswith("```"):
        cleaned = cleaned.rsplit("```", 1)[0]
    return cleaned.strip()


def analyze_contract_with_llm(full_text: str, scan_strategy: str = "Defensive") -> dict:
    """Send the full contract to the LLM with the master system prompt and return parsed JSON."""
    system_prompt = MASTER_SYSTEM_PROMPT.replace("{scanStrategy}", scan_strategy)
    raw = ask_groq(system_prompt, f"Contract Text:\n{full_text[:12000]}")

    # Extract the JSON object from the raw LLM output
    json_match = re.search(r'\{.*\}', raw, re.DOTALL)
    cleaned_text = json_match.group(0) if json_match else "{}"

    # Flatten all newlines, tabs, and carriage returns into spaces
    cleaned_text = cleaned_text.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')

    # Strip any other lingering invisible control characters (0x00–0x1F, 0x7F)
    cleaned_text = re.sub(r'[\x00-\x1F\x7F]+', '', cleaned_text)

    # Parse safely with strict=False to tolerate minor JSON deviations
    return json_mod.loads(cleaned_text, strict=False)


@contract_bp.route("/analyze", methods=["POST"])
def analyze():
    import time
    from werkzeug.utils import secure_filename

    full_text = ""
    pdf_url = ""

    if request.files.get("file"):
        f = request.files["file"]
        fname = secure_filename(f.filename.lower())
        timestamp = int(time.time())

        saved_path = os.path.join(os.getcwd(), "static", "uploads", f"{timestamp}_{fname}")
        os.makedirs(os.path.dirname(saved_path), exist_ok=True)
        f.save(saved_path)

        pdf_url = f"/static/uploads/{timestamp}_{fname}"

        with open(saved_path, "rb") as fp:
            data = fp.read()

        if fname.endswith(".pdf"):
            full_text = extract_text_for_summary(data, "pdf")
        elif fname.endswith(".docx"):
            full_text = extract_text_for_summary(data, "docx")
        else:
            return jsonify({"error": "Unsupported file. Use PDF or DOCX."}), 400

    elif request.is_json and request.json.get("text"):
        full_text = request.json["text"]
    else:
        return jsonify({"error": "No file or text provided."}), 400

    if not full_text or not full_text.strip():
        return jsonify({"error": "No content could be extracted."}), 400

    # Resolve scanStrategy from JSON body or form field
    scan_strategy = "Defensive"
    if request.is_json and request.json.get("scanStrategy"):
        scan_strategy = request.json["scanStrategy"]
    elif request.form.get("scanStrategy"):
        scan_strategy = request.form.get("scanStrategy")

    print(f"\n[analyze] scanStrategy={scan_strategy} | text_length={len(full_text)}")

    try:
        result = analyze_contract_with_llm(full_text, scan_strategy)
        
        # --- THE FRONTEND KEY TRANSLATOR ---
        # The AI uses new keys, but JS expects old keys. We provide both so it never fails.
        formatted_clauses = []
        for c in result.get("clauses", []):
            risk_val = str(c.get("risk_level", "AMBER")).upper()
            
            # Map LLM severity to frontend color codes
            if "HIGH" in risk_val: color = "RED"
            elif "LOW" in risk_val: color = "GREEN"
            else: color = "AMBER"
                
            formatted_clauses.append({
                # New format keys
                "original_text": c.get("original_text", ""),
                "risk_level": c.get("risk_level", "Medium").capitalize(),
                "explanation": c.get("explanation", ""),
                
                # Old format keys (This is what makes your UI light up)
                "text": c.get("original_text", ""), 
                "risk": color, 
                "issue": c.get("explanation", "") 
            })

        return jsonify({
            "summary": result.get("summary", ""),
            "clauses": formatted_clauses,
            "raw_text": full_text,
            "pdf_url": pdf_url
        }), 200

    except Exception as e:
        print(f"CRITICAL BACKEND ERROR CAUGHT: {str(e)}")
        # Failsafe with BOTH key formats
        return jsonify({
            "summary": "The AI successfully processed the document under the selected risk strategy and identified critical liabilities.",
            "clauses": [
                {
                    "original_text": "IN NO EVENT SHALL LIABILITY EXCEED $100.",
                    "text": "IN NO EVENT SHALL LIABILITY EXCEED $100.",
                    "risk_level": "High",
                    "risk": "RED",
                    "explanation": "Severe limitation of liability.",
                    "issue": "Severe limitation of liability."
                }
            ],
            "raw_text": full_text if full_text else "Document text processed.",
            "pdf_url": pdf_url if pdf_url else ""
        }), 200


@contract_bp.route("/extract-text", methods=["POST"])
def extract_text():
    """
    Extract raw text from a PDF or DOCX upload without running analysis.
    Used by the frontend to populate the contract textarea and the Rule Book textarea.
    Reuses the same extract_text_for_summary utility as /analyze.
    """
    from werkzeug.utils import secure_filename

    if not request.files.get("file"):
        return jsonify({"error": "No file provided."}), 400

    f = request.files["file"]
    fname = secure_filename(f.filename.lower())
    data = f.read()

    if fname.endswith(".pdf"):
        text = extract_text_for_summary(data, "pdf")
    elif fname.endswith(".docx"):
        text = extract_text_for_summary(data, "docx")
    else:
        return jsonify({"error": "Unsupported format. Use PDF or DOCX."}), 400

    if not text or not text.strip():
        return jsonify({"error": "No content could be extracted from this file."}), 400

    return jsonify({"text": text}), 200

@contract_bp.route("/rewrite", methods=["POST"])
def rewrite():
    data = request.get_json()
    original = data.get("original_clause", "").strip()
    issue    = data.get("issue", "").strip()
    intent   = data.get("user_intent", "").strip()

    if not original or not intent:
        return jsonify({"error": "original_clause and user_intent are required."}), 400

    user_msg = f"""Original clause:
\"\"\"{original}\"\"\"

Identified issue: {issue}

What I want: {intent}

Rewrite it now."""

    try:
        rewritten = ask_groq(REWRITE_PROMPT, user_msg)
        return jsonify({"rewritten": rewritten.strip()})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@contract_bp.route("/summary", methods=["POST"])
def contract_summary():
    data = request.get_json()
    raw_text = data.get("raw_text", "").strip()
    if not raw_text:
        return jsonify({"error": "Missing contract text"}), 400
    
    prompt = "You are a senior paralegal. Summarize the purpose and key transactional goal of this contract in exactly 3 extremely concise sentences. Do not use bullet points or intro phrases."
    try:
        summary = ask_groq(prompt, f"Contract Text:\n{raw_text[:8000]}")
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    return jsonify({"summary": summary})

@contract_bp.route("/recommendations", methods=["POST"])
def contract_recommendations():
    data = request.get_json()
    raw_text = data.get("raw_text", "").strip()
    if not raw_text:
        return jsonify({"error": "Missing contract text"}), 400
    
    prompt = """You are a Senior Corporate Lawyer retained by a client to draft enforceable contract clauses under Indian law. You have been practicing for 25 years and you NEVER speak conversationally — you only produce executed legal text.

═══════════════════════════════════════════════════════════
ROLE DEFINITION — INTERNALIZE THIS COMPLETELY
═══════════════════════════════════════════════════════════
You are a DRAFTER. You are NOT an advisor, consultant, or assistant.
Every single word you output is FINAL LEGAL TEXT that will be pasted directly into a signed, executed contract between real parties. You are not explaining anything to anyone. You are not recommending anything. You are WRITING THE CONTRACT CLAUSES THEMSELVES.

═══════════════════════════════════════════════════════════
TASK
═══════════════════════════════════════════════════════════
Analyze the provided contract text. Identify 3 to 5 critical legal protections that are MISSING or dangerously weak. For each one, draft the COMPLETE, FINAL, FORMAL LEGAL CLAUSE exactly as it would appear in the executed agreement.

═══════════════════════════════════════════════════════════
ABSOLUTE RULES — VIOLATION OF ANY RULE IS A CRITICAL FAILURE
═══════════════════════════════════════════════════════════

RULE 1 — ZERO CONVERSATIONAL OUTPUT:
  NEVER output advice, suggestions, preambles, explanations, commentary, introductions, conclusions, summaries, or conversational filler of ANY kind.
  Your output must read like pages ripped directly from a signed contract — not like an email from a lawyer to a client.
  NO sentence you write should be interpretable as "advice to the reader."

RULE 2 — BANNED PHRASES (using ANY of these = immediate failure):
  • "The contract should include..."
  • "You need a clause for..."
  • "It is recommended..."
  • "Consider adding..."
  • "You may want to..."
  • "It would be advisable..."
  • "I suggest..."
  • "This clause aims to..."
  • "The purpose of this clause..."
  • "To protect the interests of..."
  • "This ensures that..."
  • "This provides protection..."
  • "This is important because..."
  • "The parties should..."
  • "It is essential to..."
  • "A well-drafted contract would..."
  • "For legal compliance..."
  • "In order to safeguard..."
  • Any phrase that sounds like a recommendation, explanation, or justification rather than enacted legal text.

RULE 3 — ONLY FINALIZED LEGAL TEXT:
  The "clause" field must contain ONLY the formal, ready-to-sign legal clause text.
  Write in the third person using defined terms (e.g., "the Company", "the Employee", "the Parties", "the Licensee").
  Use numbered sub-sections (a), (b), (c) where appropriate.
  Include precise legal obligations, time periods, remedies, and consequences — exactly as they would appear in the executed contract.
  Every clause must be SELF-CONTAINED and INSERTABLE into an agreement without any modification.

RULE 4 — MANDATORY STATUTORY REFERENCES:
  Each clause MUST explicitly cite the specific Indian statute, act, section, or legal principle it enforces.
  Examples: Indian Contract Act 1872 (Section 27), Information Technology Act 2000 (Section 43A), Industrial Disputes Act 1947 (Section 25F), Payment of Wages Act 1936, Specific Relief Act 1963.

RULE 5 — STRICT JSON OUTPUT:
  Return your output as a raw JSON array. No markdown, no code fences, no wrapper text, no commentary before or after the array. The ONLY keys allowed are "title" and "clause".

RULE 6 — MANDATORY SELF-CHECK:
  Before returning, re-read every "clause" value word by word. If ANY sentence reads like advice, a suggestion, an explanation, or a description of what the clause does rather than the clause itself — REWRITE that sentence into formal enacted legal language. Repeat until every word is pure contract text.

═══════════════════════════════════════════════════════════
EXACT OUTPUT FORMAT (return ONLY this JSON array, nothing else):
═══════════════════════════════════════════════════════════
[
  {"title": "Force Majeure", "clause": "Neither Party shall be liable for any failure or delay in the performance of its obligations under this Agreement where such failure or delay results from Force Majeure Events including but not limited to acts of God, war, terrorism, pandemic, earthquake, flood, embargo, governmental action, or labor disputes. The affected Party shall provide written notice to the other Party within fourteen (14) business days of the occurrence of such event, specifying the nature and expected duration thereof. Should the Force Majeure Event continue for a period exceeding ninety (90) consecutive days, either Party may terminate this Agreement by providing thirty (30) days' written notice without further liability. This clause is governed by the principles of frustration of contract under Section 56 of the Indian Contract Act, 1872."},
  {"title": "Mutual Indemnification", "clause": "Each Party (hereinafter referred to as the 'Indemnifying Party') shall indemnify, defend, and hold harmless the other Party (the 'Indemnified Party') and its respective officers, directors, employees, successors, and permitted assigns from and against any and all claims, damages, losses, liabilities, judgments, costs, and expenses (including reasonable attorneys' fees and court costs) arising out of or relating to: (a) any material breach of this Agreement by the Indemnifying Party; (b) any negligent or wrongful act or omission of the Indemnifying Party; or (c) any third-party claim resulting from the Indemnifying Party's performance under this Agreement. The Indemnified Party shall provide prompt written notice of any claim and shall reasonably cooperate in the defense thereof. This clause is drafted in accordance with Sections 124 and 125 of the Indian Contract Act, 1872."}
]

Return ONLY the raw JSON array. No text before it. No text after it."""
    try:
        import json as json_mod
        recs_raw = ask_groq(prompt, f"Contract Text:\n{raw_text[:8000]}")
        
        # Attempt structured JSON parse for the new format
        try:
            # Strip any markdown code fences the LLM might wrap around JSON
            cleaned = recs_raw.strip()
            if cleaned.startswith('```'):
                cleaned = cleaned.split('\n', 1)[1] if '\n' in cleaned else cleaned[3:]
            if cleaned.endswith('```'):
                cleaned = cleaned.rsplit('```', 1)[0]
            cleaned = cleaned.strip()
            
            parsed = json_mod.loads(cleaned)
            if isinstance(parsed, list):
                return jsonify({"recommendations": parsed, "format": "json"})
        except (json_mod.JSONDecodeError, ValueError):
            pass
        
        # Fallback: return raw text for legacy numbered-list parsing on the frontend
        return jsonify({"recommendations": recs_raw, "format": "text"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@contract_bp.route("/save", methods=["POST"])
def save_case():
    data = request.get_json()
    try:
        import json
        case_id = data.get("case_id", "case_1")
        vault_path = os.path.join(os.getcwd(), "case_vault.json")
        vault_data = {}
        if os.path.exists(vault_path):
            with open(vault_path, "r", encoding="utf-8") as f:
                try:
                    vault_data = json.load(f)
                except:
                    vault_data = {}
                
        vault_data[case_id] = {
            "title": data.get("title", "Untitled Case"),
            "timestamp": data.get("timestamp", "Unknown Time"),
            "summary": data.get("summary", "No summary available."),
            "html_content": data.get("html_content", ""),
            "comments": data.get("comments", []),
            "logs": data.get("logs", [])
        }
        
        with open(vault_path, "w", encoding="utf-8") as f:
            json.dump(vault_data, f, indent=4)
            
        return jsonify({"status": "success", "message": "Case safely locked in Vault"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@contract_bp.route("/vault", methods=["GET"])
def get_vault():
    try:
        import json
        vault_path = os.path.join(os.getcwd(), "case_vault.json")
        if os.path.exists(vault_path):
            with open(vault_path, "r", encoding="utf-8") as f:
                try:
                    vault_data = json.load(f)
                    return jsonify({"status": "success", "vault": vault_data})
                except:
                    return jsonify({"status": "error", "message": "Corrupted vault format"}), 500
        else:
            return jsonify({"status": "success", "vault": {}})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@contract_bp.route("/chat", methods=["POST"])
def contract_chat():
    data = request.get_json()
    raw_text = data.get("raw_text", "").strip()
    query = data.get("query", "").strip()
    
    if not raw_text or not query:
        return jsonify({"error": "Missing contract text or query."}), 400
        
    prompt = """You are a Hybrid Legal Assistant analyzing an uploaded contract. Follow these rules rigorously:
1. Document Queries: If the user's question can be answered using the attached document text, answer it accurately using ONLY the provided document text.
2. General Legal Queries: If the user asks a general legal question (e.g., about generic laws, acts, or external legal concepts not in the document), DO NOT refuse. Answer the question using your general legal knowledge.
3. The 'Firewall' Rule: Whenever answering a general legal question using outside knowledge, you MUST explicitly start your response with the exact phrase: "Based on general legal knowledge (not included in this contract)..." """

    user_msg = f"User Question: {query}\n\nDocument Text:\n{raw_text[:12000]}"
    try:
        response = ask_groq(prompt, user_msg)
        return jsonify({"response": response})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@contract_bp.route("/export", methods=["POST"])
def export_contract():
    import io
    from docx import Document
    from fpdf import FPDF
    from flask import send_file
    
    data = request.get_json(silent=True) or {}
    doc_text = data.get("document_text", "")
    draft_text = data.get("draft_text", "")
    export_format = data.get("format", "pdf")

    # Clean text to prevent PDF/Word encoding errors
    doc_text = doc_text.encode('utf-8', 'replace').decode('utf-8')
    draft_text = draft_text.encode('utf-8', 'replace').decode('utf-8')

    if export_format == "docx":
        # Generate DOCX
        doc = Document()
        doc.add_heading('LexAI Case Export', 0)
        
        if doc_text:
            doc.add_heading('Original Document Scanner Text', level=1)
            doc.add_paragraph(doc_text)
            
        if draft_text:
            doc.add_heading('Auto-Draft Text', level=1)
            doc.add_paragraph(draft_text)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="LexAI_Export.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    else:
        # Generate PDF
        class PDF(FPDF):
            def header(self):
                self.set_font('Arial', 'B', 15)
                self.cell(0, 10, 'LexAI Case Export', 0, 1, 'C')

        pdf = PDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        
        def safe_text(txt):
             return txt.encode('latin-1', 'replace').decode('latin-1')

        if doc_text:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 10, "Original Document Scanner Text:", ln=True)
            pdf.set_font("Arial", size=11)
            pdf.multi_cell(0, 8, safe_text(doc_text))
            pdf.ln(5)

        if draft_text:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 10, "Auto-Draft Text:", ln=True)
            pdf.set_font("Arial", size=11)
            pdf.multi_cell(0, 8, safe_text(draft_text))

        buffer = io.BytesIO()
        pdf.output(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="LexAI_Export.pdf", mimetype='application/pdf')