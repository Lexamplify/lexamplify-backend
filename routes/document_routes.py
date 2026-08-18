"""
routes/document_routes.py
Blueprint: /api/documents
  POST /api/documents/draft        — AI clause/document synthesis (instructions + optional precedent/context)
  POST /api/documents/inline-edit  — AI rewrite of a selected editor range (BubbleMenu "AI Rewrite")
  POST /api/documents/upload       — Upload files (PDF/DOCX/TXT), extract text, save and vectorize (RAG)
  GET  /api/documents              — List uploaded documents (filtered by case_id)
  DELETE /api/documents/<id>       — Delete document and cascade delete RAG vectors
"""
import os
import io
import re
import uuid
import sqlite3
from flask import Blueprint, jsonify, request, current_app
from flask_jwt_extended import jwt_required, get_jwt_identity
from werkzeug.utils import secure_filename
from utils.ai_helper import ask_groq
from utils.rag_pipeline import ingest_document

doc_bp = Blueprint("document", __name__)

# Strips a ```html / ``` code fence and stray wrapping quotes the LLM adds
# despite being told not to — same defensive posture as
# extract_json_from_llm_response in ai_helper.py, applied to HTML instead
# of JSON.
_CODE_FENCE_RE = re.compile(r'^```(?:html)?\s*|\s*```$', re.MULTILINE)


def _clean_inline_html(raw: str) -> str:
    text = _CODE_FENCE_RE.sub('', raw).strip()
    if len(text) >= 2 and text[0] == '"' and text[-1] == '"':
        text = text[1:-1].strip()
    return text

DB_PATH = "lex_assistant.db"

# ── 1. ROBUST TEXT EXTRACTION UTILITY ───────────────────────────────────

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extracts plain text from PDF, DOCX, or TXT file using fitz (PyMuPDF), pdfplumber,
    python-docx, or PyPDF2 with cascading fallbacks.
    """
    ext = filename.lower().split('.')[-1]
    
    if ext == 'pdf':
        # Fallback Level 1: PyMuPDF (fitz)
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            if text.strip():
                return text.strip()
        except ImportError:
            print("[PDF Extractor] PyMuPDF (fitz) not installed. Trying pdfplumber...")
        except Exception as e:
            print(f"[PDF Extractor] PyMuPDF extraction failed: {e}")
            
        # Fallback Level 2: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if text.strip():
                    return text.strip()
        except ImportError:
            print("[PDF Extractor] pdfplumber not installed. Trying PyPDF2...")
        except Exception as e:
            print(f"[PDF Extractor] pdfplumber extraction failed: {e}")

        # Fallback Level 3: PyPDF2 (guaranteed standard)
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
            return text.strip()
        except Exception as e:
            print(f"[PDF Extractor] PyPDF2 fallback failed: {e}")
            return ""

    elif ext in ['docx', 'doc']:
        # Extract text and tables from DOCX
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            content_list = []
            
            # Paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    content_list.append(p.text)
                    
            # Table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text and text not in content_list:
                            content_list.append(text)
                            
            return "\n".join(content_list).strip()
        except Exception as e:
            print(f"[DOCX Extractor] python-docx extraction failed: {e}")
            return ""

    else:
        # Plain text
        try:
            return file_bytes.decode('utf-8', errors='ignore').strip()
        except Exception as e:
            print(f"[Text Extractor] plain text decode failed: {e}")
            return ""

# ── 2. AUTO-DRAFT (AI CLAUSE SYNTHESIS) ROUTE ──────────────────────────
# Replaces an earlier version of this same route that used a different
# request contract (prompt/context, required both, returned a bare
# {"draft": ...} with no "status" key) — that shape is preserved as a
# fallback in the payload parsing below so ContractAnalyzer.jsx's existing
# Auto-Draft panel (which still posts {prompt, context}) keeps working
# unchanged.

@doc_bp.route("/draft", methods=["POST", "OPTIONS"])
def auto_draft():
    if request.method == "OPTIONS":
        # No manual CORS headers here — app.py's global @app.after_request
        # hook already adds Access-Control-Allow-Origin as an
        # origin-validated echo (never "*"; that hook's own comment
        # explains why: Allow-Credentials requires a specific origin, not
        # a wildcard). Setting "*" here would either get silently
        # overwritten for validated origins or leak an open CORS response
        # to anything that ISN'T validated — matching the plain
        # `return jsonify({}), 200` every other OPTIONS handler in this
        # codebase already uses.
        return jsonify({}), 200

    data = request.get_json(force=True, silent=True) or {}
    # str(...) around each fallback chain guards against a non-string JSON
    # value (e.g. {"instructions": 12345}) reaching .strip() below and
    # crashing with an unhandled AttributeError instead of the intended
    # 400 — "bulletproof" parsing means surviving malformed types, not
    # just missing keys. .strip() applied at assignment so every use below
    # reads the already-normalized value, not a mix of stripped/unstripped.
    instructions = str(data.get('instructions') or data.get('drafting_instructions') or data.get('prompt') or '').strip()
    precedent = str(data.get('precedent_insert') or data.get('precedent') or '').strip()
    # 'context' (no prefix) is the field name the existing
    # ContractAnalyzer.jsx caller already sends — kept as a fallback
    # alongside the new 'reference_context' name so that live integration
    # doesn't silently lose its context on this refactor.
    context = str(data.get('reference_context') or data.get('context') or '').strip()

    if not instructions:
        return jsonify({"error": True, "message": "Drafting instructions required."}), 400

    depth = str(data.get('depth') or 'comprehensive').strip()

    system_prompt = (
        "You are a Senior Indian Legal Advocate & Corporate Law Partner. "
        "Synthesize a highly detailed, comprehensive, enterprise-grade legal clause or agreement draft based strictly on Indian Law.\n\n"
        "FORMATTING & STRUCTURE REQUIREMENTS:\n"
        "1. Use Markdown formatting: Use level-3 headings (### 1. Title of Clause) for main headings.\n"
        "2. Format clauses into distinct numbered sub-clauses (e.g. 1.1, 1.2, 1.3) with double line breaks between paragraphs so that clauses are cleanly spaced.\n"
        "3. For sub-conditions or itemized lists, use lettered indents (e.g. (a), (b), (c)) on separate lines.\n"
        "4. Include explicit Indian statutory citations (e.g., **Indian Contract Act, 1872**, **Arbitration & Conciliation Act, 1996**, **Copyright Act, 1957**, **Specific Relief Act, 1963**, **Information Technology Act, 2000**) wherever applicable.\n"
        "5. Depth & Details: Provide thorough, enterprise-level depth with operative obligations, notice requirements, cure periods, remedies, and governing law. Do NOT produce short 2-sentence summaries. Write a complete, execution-ready legal text without conversational fluff or preambles.\n\n"
        "MANDATORY INSTRUCTION: You must generate the complete agreement from Title, Parties, Recitals, Operative Clauses (1 through N), to Boilerplate (Severability, Notices, Jurisdiction), concluding strictly with the formal Execution & Signature Block. Never truncate, omit sections, or leave trailing markdown tokens."
    )
    if depth == 'comprehensive':
        system_prompt += "\n6. Include full definitions, operating obligations, indemnity scope, liability caps, and dispute escalation steps."
    if context:
        system_prompt += f"\n\nREFERENCE CONTEXT:\n{context}"
    if precedent:
        system_prompt += f"\n\nPRECEDENT TO INCORPORATE:\n{precedent}"

    try:
        generated_text = ask_groq(system_prompt, f"Drafting instructions: {instructions}", max_tokens=8192, timeout=120)
        if not generated_text or not generated_text.strip():
            raise ValueError("LLM returned an empty draft.")
        generated_text = generated_text.strip()
        # draft/clause/content are the same string under three names —
        # different callers (old and new UI panels) read different keys
        # for the identical generated text, so all three ship together
        # rather than requiring each caller to agree on one key first.
        return jsonify({
            "status": "success",
            "draft": generated_text,
            "clause": generated_text,
            "content": generated_text,
        }), 200
    except Exception as e:
        print(f"[Auto-Draft Error]: {e}")
        return jsonify({"error": True, "message": "AI reasoning engine timeout or failure. Please retry."}), 500

# ── 2b. INLINE AI REWRITE (BubbleMenu "AI Rewrite") ────────────────────
# Rewrites an arbitrary selection from the contract editor in place. Unlike
# /draft (which synthesizes a whole new clause), the caller here already
# has a locked {from, to} range in the live ProseMirror doc and replaces
# it directly with whatever HTML comes back — so the output MUST be a
# small inline-safe fragment, not a full document.

@doc_bp.route("/inline-edit", methods=["POST", "OPTIONS"])
def inline_edit():
    if request.method == "OPTIONS":
        return jsonify({}), 200

    data = request.get_json(force=True, silent=True) or {}
    selected_text = str(data.get('selectedText') or data.get('selected_text') or '').strip()
    instruction = str(data.get('instruction') or '').strip()

    if not selected_text or not instruction:
        return jsonify({"error": True, "message": "selectedText and instruction are required."}), 400

    system_prompt = (
        "You are an inline legal drafting assistant embedded in a rich-text contract editor. "
        "Rewrite the SELECTED TEXT by applying the INSTRUCTION, preserving legal meaning and Indian "
        "statutory terminology unless the instruction says otherwise.\n\n"
        "Output rules (strict):\n"
        "1. Output ONLY the rewritten result as raw HTML — no Markdown code fences (no ```), no "
        "explanation, no preamble, no surrounding quotes.\n"
        "2. The result replaces a selection INSIDE an existing paragraph — use inline formatting only "
        "(<strong>, <em>, <u>, <br>). Do NOT wrap the output in <p>, <div>, <html>, or <body> tags.\n"
        "3. If no formatting is needed, return plain text with no tags at all."
    )
    user_msg = f"SELECTED TEXT:\n{selected_text}\n\nINSTRUCTION: {instruction}"

    try:
        raw = ask_groq(system_prompt, user_msg)
        if not raw or not raw.strip():
            raise ValueError("LLM returned an empty rewrite.")
        html = _clean_inline_html(raw)
        if not html:
            raise ValueError("LLM rewrite was empty after cleanup.")
        return jsonify({"status": "success", "html": html}), 200
    except Exception as e:
        print(f"[Inline Edit Error]: {e}")
        return jsonify({"error": True, "message": "AI inline rewrite failed. Please retry."}), 500

# ── 3. ENTERPRISE UPLOAD & RAG INGESTION ROUTE ─────────────────────────

@doc_bp.route("/upload", methods=["POST"])
@jwt_required()
def upload_document():
    """
    Accepts document files, executes clean text extraction, generates a two-sentence
    summary, and indexes the document chunks with embeddings in the database.
    """
    user_id = int(get_jwt_identity())
    
    if 'file' not in request.files:
        return jsonify({"error": "No file part in the multipart request."}), 400
        
    f = request.files['file']
    if f.filename == '':
        return jsonify({"error": "No file selected."}), 400
        
    # case_vault.case_id is TEXT (mass-ingestion and auto-provisioned vault
    # entries use non-numeric ids like "case_138_dashrath" or
    # "conflict-<timestamp>"), so keep it as a plain string rather than
    # forcing int() the way the old SQLAlchemy FK column required.
    case_id = request.form.get("case_id") or None

    try:
        # Read file bytes in memory for extraction
        file_bytes = f.read()
        extracted_text = extract_text_from_file(file_bytes, f.filename)
        
        if not extracted_text.strip():
            return jsonify({"error": "Failed to extract clean text from the document. The file might be scanned/empty."}), 400
            
        # Save file to uploads folder
        upload_folder = current_app.config.get('UPLOAD_FOLDER', 'static/uploads')
        os.makedirs(upload_folder, exist_ok=True)
        unique_filename = f"{uuid.uuid4().hex}_{secure_filename(f.filename)}"
        file_path = os.path.join(upload_folder, unique_filename)
        
        with open(file_path, "wb") as out_file:
            out_file.write(file_bytes)
            
        # Generate a lightweight 2-sentence summary
        summary = "No summary generated."
        try:
            summary_system_prompt = (
                "You are an elite Indian legal assistant. Read the provided text and write a brief, "
                "2-sentence executive summary of the document highlighting the parties, type of document, "
                "and key subject matter. Keep it strictly to two sentences."
            )
            # Sample first 4000 characters to make summary generation fast
            summary = ask_groq(summary_system_prompt, f"Text snippet:\n{extracted_text[:4000]}").strip()
        except Exception as se:
            print(f"[RAG Ingestion] AI Summary generation skipped: {se}")
            summary = f"Uploaded legal document '{f.filename}' containing {len(extracted_text)} characters."

        # Save metadata record into case_vault — the single source of truth
        # shared with Firm Library / LexAmplify, instead of the old SQLAlchemy
        # Document table. doc_type defaults to "Vault Document" and tags to
        # a valid empty JSON array so the citation json_insert/json_remove
        # SQL functions never choke on a NULL or malformed value.
        conn = sqlite3.connect(DB_PATH)
        try:
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO case_vault (case_id, title, doc_type, content, tags) VALUES (?, ?, ?, ?, ?)",
                (case_id, f.filename, "Vault Document", extracted_text, "[]")
            )
            conn.commit()
            new_doc_id = cursor.lastrowid
        finally:
            conn.close()

        # Ingest text chunks & generate embeddings — keyed to the new
        # case_vault id so RAG chunk lookups stay consistent with it.
        chunks_count = ingest_document(
            document_id=new_doc_id,
            case_id=case_id,
            user_id=user_id,
            text=extracted_text
        )

        return jsonify({
            "message": "Document uploaded and vectorized successfully.",
            "document": {
                "id": new_doc_id,
                "filename": f.filename,
                "summary": summary,
                "chunks_indexed": chunks_count
            }
        }), 201

    except Exception as e:
        print(f"[Upload Route Error]: {e}")
        return jsonify({"error": f"Failed to ingest document: {str(e)}"}), 500

# ── 4. DOCUMENT LIST ROUTE ──────────────────────────────────────────────

@doc_bp.route("", methods=["GET"])
@jwt_required()
def list_documents():
    """Lists metadata for case_vault documents, optionally filtered by
    case_id. Queries case_vault directly — the single source of truth shared
    with Firm Library / LexAmplify — instead of the old SQLAlchemy Document table,
    so ids returned here line up with the citation routes' doc_id space."""
    case_id_raw = request.args.get("case_id")

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        if case_id_raw:
            rows = conn.execute(
                "SELECT id AS doc_id, case_id, title, doc_type, tags, created_at "
                "FROM case_vault WHERE case_id = ? ORDER BY created_at DESC",
                (case_id_raw,)
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT id AS doc_id, case_id, title, doc_type, tags, created_at "
                "FROM case_vault ORDER BY created_at DESC"
            ).fetchall()
    except Exception as e:
        print(f"[List Docs Error]: {e}")
        return jsonify({"error": "Failed to fetch documents."}), 500
    finally:
        conn.close()

    # Mapped onto the exact shape CaseVault.jsx already renders (id/filename/
    # filetype/summary) so the existing table UI needs no changes.
    return jsonify([{
        "id": r["doc_id"],
        "case_id": r["case_id"],
        "filename": r["title"],
        "filetype": r["doc_type"],
        "summary": None,
        "tags": r["tags"],
        "created_at": r["created_at"],
    } for r in rows]), 200

# ── 5. DOCUMENT DETAILS ROUTE (WITH CHUNKS RECONSTRUCTION) ──────────────

@doc_bp.route("/<int:doc_id>", methods=["GET"])
@jwt_required()
def get_document_details(doc_id):
    """Fetches case_vault document metadata and reconstructs full text from
    its RAG chunks, falling back to the row's own stored content if no
    chunks exist."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        row = conn.execute(
            "SELECT id AS doc_id, case_id, title, doc_type, tags, content, created_at "
            "FROM case_vault WHERE id = ?",
            (doc_id,)
        ).fetchone()
        if not row:
            return jsonify({"error": "Document not found."}), 404

        chunk_rows = conn.execute(
            "SELECT chunk_text FROM document_chunks WHERE document_id = ? ORDER BY chunk_index ASC",
            (doc_id,)
        ).fetchall()
    except Exception as e:
        print(f"[Get Doc Details Error]: {e}")
        return jsonify({"error": f"Failed to retrieve document details: {str(e)}"}), 500
    finally:
        conn.close()

    full_text = "\n\n".join(r[0] for r in chunk_rows) if chunk_rows else (row["content"] or "")

    return jsonify({
        "id": row["doc_id"],
        "case_id": row["case_id"],
        "filename": row["title"],
        "filetype": row["doc_type"],
        "summary": None,
        "tags": row["tags"],
        "created_at": row["created_at"],
        "text": full_text
    }), 200

# ── 6. DOCUMENT DELETE ROUTE ───────────────────────────────────────────

@doc_bp.route("/<int:doc_id>", methods=["DELETE"])
@jwt_required()
def delete_document(doc_id):
    """Deletes the case_vault document and cascade-deletes its RAG chunks."""
    conn = sqlite3.connect(DB_PATH)
    try:
        row = conn.execute("SELECT id FROM case_vault WHERE id = ?", (doc_id,)).fetchone()
        if not row:
            return jsonify({"error": "Document not found."}), 404

        conn.execute("DELETE FROM document_chunks WHERE document_id = ?", (doc_id,))
        conn.execute("DELETE FROM case_vault WHERE id = ?", (doc_id,))
        conn.commit()
        return jsonify({"message": "Document and vectorized index successfully deleted."}), 200
    except Exception as e:
        conn.rollback()
        print(f"[Delete Doc Error]: {e}")
        return jsonify({"error": f"Failed to delete document: {str(e)}"}), 500
    finally:
        conn.close()
